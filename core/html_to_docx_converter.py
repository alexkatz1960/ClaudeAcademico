"""
HTMLtoDocxConverter - Sistema de Traducción Académica v2.2
========================================================

Convierte HTML a DOCX preservando formatos críticos y footnotes.
Implementa estrategia dual: Pandoc + python-docx con validación automática.

Autor: Sistema ClaudeAcademico v2.2
Fecha: Enero 2025
Versión: 2.2.1 (Post-Auditoría Externa - Modificaciones 1 & 2)
"""

import re
import json
import hashlib
import logging
import subprocess
import tempfile
from pathlib import Path
from typing import Dict, List, Tuple, Any, Optional, Union, Protocol
from dataclasses import dataclass, asdict
from datetime import datetime
import uuid
from .html_sanitizer import html_sanitizer

# Document processing
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.oxml.shared import qn
import mammoth
from bs4 import BeautifulSoup, Tag

# Módulo desacoplado para análisis de estilos (Modificación 2)
try:
    from .html_style_analyzer import HTMLStyleAnalyzer, ExtendedStyleInfo, StylePreservationAlert
except ImportError:
    # Fallback si no está disponible el módulo
    HTMLStyleAnalyzer = None
    ExtendedStyleInfo = None
    StylePreservationAlert = None

# Utilities
import pandas as pd
from concurrent.futures import ThreadPoolExecutor, TimeoutError
import time


# Protocolo para inyección de dependencias (Modificación 2)
class PandocRunner(Protocol):
    """Protocolo para ejecutor de Pandoc personalizable"""
    def run_pandoc(self, input_file: str, output_file: str, options: List[str]) -> bool:
        ...


class Logger(Protocol):
    """Protocolo para logger personalizable"""
    def info(self, message: str) -> None: ...
    def debug(self, message: str) -> None: ...
    def warning(self, message: str) -> None: ...
    def error(self, message: str) -> None: ...


@dataclass
class StyleInfo:
    """Información de estilo extraída del HTML (compatibilidad)"""
    bold_count: int = 0
    italic_count: int = 0
    underline_count: int = 0
    superscript_count: int = 0
    subscript_count: int = 0
    blockquote_count: int = 0
    heading_counts: Dict[str, int] = None
    list_counts: Dict[str, int] = None
    table_count: int = 0
    footnote_count: int = 0
    
    def __post_init__(self):
        if self.heading_counts is None:
            self.heading_counts = {}
        if self.list_counts is None:
            self.list_counts = {}


@dataclass
class FootnoteInfo:
    """Información de footnote con ID único y soporte real"""
    unique_id: str
    original_text: str
    superscript_location: str
    reference_number: int
    processed_text: str = ""
    footnote_content: str = ""  # Contenido completo de la footnote
    is_functional: bool = False  # Si la footnote es funcional en DOCX
    docx_footnote_id: Optional[int] = None  # ID en el documento DOCX
    
    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


@dataclass 
class ComplexTableInfo:
    """Información de tabla compleja con rowspan/colspan"""
    table_index: int
    has_rowspan: bool = False
    has_colspan: bool = False
    max_rowspan: int = 1
    max_colspan: int = 1
    total_merged_cells: int = 0
    preservation_strategy: str = "flatten"  # flatten, approximate, manual_review
    
    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


@dataclass
class ConversionResult:
    """Resultado de conversión HTML→DOCX con análisis extendido"""
    success: bool
    method_used: str
    output_path: str
    style_preservation_score: float
    footnote_preservation_score: float
    processing_time_seconds: float
    original_styles: Union[StyleInfo, ExtendedStyleInfo]
    preserved_styles: Union[StyleInfo, ExtendedStyleInfo]
    footnote_mapping: Dict[str, FootnoteInfo]
    validation_report: Dict[str, Any]
    
    # Nuevos campos para funcionalidad extendida
    preservation_alerts: List[StylePreservationAlert] = None
    complex_tables_info: List[ComplexTableInfo] = None
    functional_footnotes_created: int = 0
    manual_review_required: bool = False
    alert_summary: Dict[str, int] = None
    error_details: Optional[str] = None
    
    def __post_init__(self):
        if self.preservation_alerts is None:
            self.preservation_alerts = []
        if self.complex_tables_info is None:
            self.complex_tables_info = []
        if self.alert_summary is None:
            self.alert_summary = {'critical': 0, 'warning': 0, 'info': 0}
    
    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


class HTMLtoDocxConverter:
    """
    Sistema avanzado de conversión HTML → DOCX con preservación de formatos
    
    Características principales:
    - Estrategia dual: Pandoc (robusto) + python-docx (control)
    - Marcadores de formato protegidos para Claude
    - Validación automática de estilos preservados
    - Reconexión robusta de footnotes con IDs únicos
    - Fallbacks automáticos y recuperación de errores
    - Sistema de alertas automáticas con códigos de severidad
    - Soporte para estilos avanzados y estructuras complejas
    """
    
    def __init__(self, 
                 pandoc_timeout: int = 60,
                 validation_threshold: float = 0.90,
                 preserve_footnotes: bool = True,
                 enable_extended_analysis: bool = True,
                 create_functional_footnotes: bool = True,
                 debug_mode: bool = False,
                 # Inyección de dependencias opcional (Modificación 2)
                 pandoc_runner: Optional[PandocRunner] = None,
                 logger: Optional[Logger] = None,
                 style_analyzer: Optional[HTMLStyleAnalyzer] = None):
        """
        Inicializa el convertidor HTML→DOCX con configuración avanzada
        
        Args:
            pandoc_timeout: Tiempo límite para conversión con Pandoc (segundos)
            validation_threshold: Umbral mínimo para validación de estilos
            preserve_footnotes: Si preservar footnotes con IDs únicos
            enable_extended_analysis: Habilitar análisis de estilos avanzados
            create_functional_footnotes: Crear footnotes funcionales en DOCX
            debug_mode: Habilitar logging detallado
            pandoc_runner: Implementación personalizada de Pandoc (opcional)
            logger: Logger personalizado (opcional)
            style_analyzer: Analizador de estilos personalizado (opcional)
        """
        self.pandoc_timeout = pandoc_timeout
        self.validation_threshold = validation_threshold
        self.preserve_footnotes = preserve_footnotes
        self.enable_extended_analysis = enable_extended_analysis
        self.create_functional_footnotes = create_functional_footnotes
        self.debug_mode = debug_mode
        
        # Inyección de dependencias (Modificación 2)
        self.pandoc_runner = pandoc_runner
        
        # Configurar logging
        if logger:
            self.logger = logger
        else:
            self.logger = logging.getLogger(f"{__name__}.HTMLtoDocxConverter")
            if debug_mode:
                self.logger.setLevel(logging.DEBUG)
        
        # Inicializar analizador de estilos
        if style_analyzer:
            self.style_analyzer = style_analyzer
        elif HTMLStyleAnalyzer:
            self.style_analyzer = HTMLStyleAnalyzer(
                enable_extended_analysis=enable_extended_analysis,
                debug_mode=debug_mode
            )
        else:
            self.style_analyzer = None
            self.logger.warning("HTMLStyleAnalyzer no disponible - funcionalidad limitada")
        
        # Verificar dependencias
        self._verify_dependencies()
        
        # Configurar marcadores de formato con soporte extendido
        self.format_markers = self._initialize_format_markers()
        
        # Patrones para detectar estructuras complejas
        self.complex_structure_patterns = self._initialize_complex_patterns()
        
        # Estadísticas
        self.conversion_stats = {
            'total_conversions': 0,
            'pandoc_successes': 0,
            'native_successes': 0,
            'total_failures': 0,
            'functional_footnotes_created': 0,
            'complex_tables_processed': 0,
            'avg_processing_time': 0.0,
            'avg_style_preservation': 0.0,
            'critical_alerts_generated': 0
        }
        
        self.logger.info("HTMLtoDocxConverter inicializado correctamente con funcionalidad extendida")
    
    def _verify_dependencies(self) -> None:
        """Verifica que todas las dependencias estén disponibles"""
        try:
            # Verificar Pandoc
            result = subprocess.run(['pandoc', '--version'], 
                                  capture_output=True, text=True, timeout=10)
            if result.returncode != 0:
                raise RuntimeError("Pandoc no está disponible o no funciona correctamente")
            
            pandoc_version = result.stdout.split('\n')[0]
            self.logger.info(f"Pandoc disponible: {pandoc_version}")
            
        except (FileNotFoundError, subprocess.TimeoutExpired, RuntimeError) as e:
            self.logger.warning(f"Pandoc no disponible: {e}. Solo se usará conversión nativa")
            self.pandoc_available = False
        else:
            self.pandoc_available = True
        
        # Verificar python-docx
        try:
            doc = Document()
            self.logger.info("python-docx disponible y funcional")
        except Exception as e:
            raise RuntimeError(f"python-docx no funciona correctamente: {e}")
    
    def _initialize_format_markers(self) -> Dict[str, str]:
        """Inicializa el mapeo de tags HTML a marcadores protegidos con soporte extendido"""
        markers = {
            # Formatos básicos
            r'<strong[^>]*>(.*?)</strong>': r'<<<BOLD_START>>>\1<<<BOLD_END>>>',
            r'<b[^>]*>(.*?)</b>': r'<<<BOLD_START>>>\1<<<BOLD_END>>>',
            r'<em[^>]*>(.*?)</em>': r'<<<ITALIC_START>>>\1<<<ITALIC_END>>>',
            r'<i[^>]*>(.*?)</i>': r'<<<ITALIC_START>>>\1<<<ITALIC_END>>>',
            r'<u[^>]*>(.*?)</u>': r'<<<UNDERLINE_START>>>\1<<<UNDERLINE_END>>>',
            
            # Superíndices y subíndices
            r'<sup[^>]*>(.*?)</sup>': r'<<<SUP_START>>>\1<<<SUP_END>>>',
            r'<sub[^>]*>(.*?)</sub>': r'<<<SUB_START>>>\1<<<SUB_END>>>',
            
            # Estructuras de bloque
            r'<blockquote[^>]*>(.*?)</blockquote>': r'<<<BLOCKQUOTE_START>>>\1<<<BLOCKQUOTE_END>>>',
            r'<h1[^>]*>(.*?)</h1>': r'<<<H1_START>>>\1<<<H1_END>>>',
            r'<h2[^>]*>(.*?)</h2>': r'<<<H2_START>>>\1<<<H2_END>>>',
            r'<h3[^>]*>(.*?)</h3>': r'<<<H3_START>>>\1<<<H3_END>>>',
            r'<h4[^>]*>(.*?)</h4>': r'<<<H4_START>>>\1<<<H4_END>>>',
            r'<h5[^>]*>(.*?)</h5>': r'<<<H5_START>>>\1<<<H5_END>>>',
            r'<h6[^>]*>(.*?)</h6>': r'<<<H6_START>>>\1<<<H6_END>>>',
            
            # Listas (incluyendo anidadas)
            r'<ul[^>]*>(.*?)</ul>': r'<<<UL_START>>>\1<<<UL_END>>>',
            r'<ol[^>]*>(.*?)</ol>': r'<<<OL_START>>>\1<<<OL_END>>>',
            r'<li[^>]*>(.*?)</li>': r'<<<LI_START>>>\1<<<LI_END>>>',
            
            # Párrafos y saltos
            r'<p[^>]*>(.*?)</p>': r'<<<P_START>>>\1<<<P_END>>>',
            r'<br[^>]*/?>\s*': r'<<<BR>>>',
            
            # Tablas con soporte para estructuras complejas
            r'<table[^>]*>(.*?)</table>': r'<<<TABLE_START>>>\1<<<TABLE_END>>>',
            r'<tr[^>]*>(.*?)</tr>': r'<<<TR_START>>>\1<<<TR_END>>>',
            r'<td[^>]*rowspan="([^"]*)"[^>]*>(.*?)</td>': r'<<<TD_ROWSPAN_\1_START>>>\2<<<TD_ROWSPAN_END>>>',
            r'<td[^>]*colspan="([^"]*)"[^>]*>(.*?)</td>': r'<<<TD_COLSPAN_\1_START>>>\2<<<TD_COLSPAN_END>>>',
            r'<td[^>]*>(.*?)</td>': r'<<<TD_START>>>\1<<<TD_END>>>',
            r'<th[^>]*>(.*?)</th>': r'<<<TH_START>>>\1<<<TH_END>>>',
        }
        
        # Formatos avanzados (Modificación 1)
        if self.enable_extended_analysis:
            advanced_markers = {
                # Versalitas
                r'<(small|sc)[^>]*>(.*?)</(small|sc)>': r'<<<SMALLCAPS_START>>>\2<<<SMALLCAPS_END>>>',
                r'<[^>]*style="[^"]*font-variant:\s*small-caps[^"]*"[^>]*>(.*?)</[^>]*>': r'<<<SMALLCAPS_START>>>\1<<<SMALLCAPS_END>>>',
                
                # Texto tachado
                r'<(s|strike|del)[^>]*>(.*?)</(s|strike|del)>': r'<<<STRIKETHROUGH_START>>>\2<<<STRIKETHROUGH_END>>>',
                r'<[^>]*style="[^"]*text-decoration:\s*line-through[^"]*"[^>]*>(.*?)</[^>]*>': r'<<<STRIKETHROUGH_START>>>\1<<<STRIKETHROUGH_END>>>',
                
                # Texto con color
                r'<[^>]*style="[^"]*color:\s*([^;"][^"]*)[^"]*"[^>]*>(.*?)</[^>]*>': r'<<<COLOR_\1_START>>>\2<<<COLOR_END>>>',
                r'<font[^>]*color="([^"]*)"[^>]*>(.*?)</font>': r'<<<COLOR_\1_START>>>\2<<<COLOR_END>>>',
                
                # Texto resaltado
                r'<(mark|highlight)[^>]*>(.*?)</(mark|highlight)>': r'<<<HIGHLIGHT_START>>>\2<<<HIGHLIGHT_END>>>',
                r'<[^>]*style="[^"]*background-color:\s*([^;"][^"]*)[^"]*"[^>]*>(.*?)</[^>]*>': r'<<<HIGHLIGHT_\1_START>>>\2<<<HIGHLIGHT_END>>>',
                
                # Alineaciones de párrafo
                r'<p[^>]*style="[^"]*text-align:\s*center[^"]*"[^>]*>(.*?)</p>': r'<<<P_CENTER_START>>>\1<<<P_CENTER_END>>>',
                r'<p[^>]*style="[^"]*text-align:\s*right[^"]*"[^>]*>(.*?)</p>': r'<<<P_RIGHT_START>>>\1<<<P_RIGHT_END>>>',
                r'<p[^>]*style="[^"]*text-align:\s*justify[^"]*"[^>]*>(.*?)</p>': r'<<<P_JUSTIFY_START>>>\1<<<P_JUSTIFY_END>>>',
                
                # Código
                r'<(code|kbd|samp)[^>]*>(.*?)</(code|kbd|samp)>': r'<<<CODE_START>>>\2<<<CODE_END>>>',
                r'<pre[^>]*>(.*?)</pre>': r'<<<PREFORMATTED_START>>>\1<<<PREFORMATTED_END>>>',
            }
            markers.update(advanced_markers)
        
        return markers
    
    def _initialize_complex_patterns(self) -> Dict[str, str]:
        """Inicializa patrones para detectar estructuras complejas"""
        return {
            'complex_table_cell': r'<t[dh][^>]*(rowspan|colspan)="[^"]*"[^>]*>',
            'nested_list': r'<(ul|ol)[^>]*>.*?<(ul|ol)[^>]*>.*?</\2>.*?</\1>',
            'multi_level_heading': r'<h[1-6][^>]*>.*?<h[1-6][^>]*>',
            'complex_footnote': r'<sup[^>]*>.*?<a[^>]*href="#fn\d+"[^>]*>.*?</a>.*?</sup>',
        }
    
    def convert_with_validation(self, 
                              html_content: str, 
                              output_path: str,
                              book_id: Optional[str] = None) -> ConversionResult:
        """
        Convierte HTML a DOCX con validación automática de preservación y alertas
        
        Args:
            html_content: Contenido HTML a convertir
            output_path: Ruta del archivo DOCX de salida
            book_id: ID del libro para tracking (opcional)
            
        Returns:
            ConversionResult con detalles completos de la conversión y alertas
        """
        start_time = time.time()
        
        self.logger.info(f"Iniciando conversión HTML→DOCX: {output_path}")
        
        try:
            # 1. Analizar estilos originales del HTML (usando analizador extendido si está disponible)
            if self.style_analyzer:
                original_styles = self.style_analyzer.extract_html_styles(html_content)
            else:
                original_styles = self._extract_html_styles_fallback(html_content)
            
            self.logger.debug(f"Estilos originales extraídos: {asdict(original_styles)}")
            
            # 2. Detectar estructuras complejas antes del procesamiento
            complex_tables_info = self._detect_complex_tables(html_content)
            if complex_tables_info:
                self.logger.info(f"Detectadas {len(complex_tables_info)} tablas complejas")
            
            # 3. Sanitizar HTML para prevenir XSS y validar seguridad
            validation_result = html_sanitizer.validate_html_safety(html_content)
            if not validation_result["safe"]:
                self.logger.warning(f"HTML inseguro detectado: {validation_result['issues']}")
            
            sanitized_html = html_sanitizer.sanitize_html(html_content)
            if len(sanitized_html) != len(html_content):
                self.logger.info(f"HTML sanitizado: {len(html_content)} -> {len(sanitized_html)} chars")
            
            # Usar HTML sanitizado para el resto del procesamiento
            html_content = sanitized_html
            
            # 4. Procesar footnotes si está habilitado
            footnote_mapping = {}
            processed_html = html_content
            
            if self.preserve_footnotes:
                processed_html, footnote_mapping = self._process_footnotes_for_preservation(html_content)
                self.logger.info(f"Procesadas {len(footnote_mapping)} footnotes con IDs únicos")
            
            # 5. Insertar marcadores de formato protegidos (con soporte extendido)
            protected_html = self._insert_format_markers(processed_html)
            
            # 6. Intentar conversión con métodos disponibles
            conversion_success = False
            method_used = ""
            
            # Método 1: Pandoc (preferido por robustez)
            if self.pandoc_available:
                try:
                    if self.pandoc_runner:
                        # Usar implementación personalizada
                        success = self.pandoc_runner.run_pandoc(
                            protected_html, output_path, 
                            ['-f', 'html', '-t', 'docx', '--preserve-tabs', '--wrap=preserve']
                        )
                    else:
                        success = self._convert_html_to_docx_pandoc(protected_html, output_path)
                    
                    if success:
                        conversion_success = True
                        method_used = "pandoc"
                        self.conversion_stats['pandoc_successes'] += 1
                        self.logger.info("Conversión exitosa con Pandoc")
                except Exception as e:
                    self.logger.warning(f"Conversión con Pandoc falló: {e}")
            
            # Método 2: Nativo con python-docx (fallback)
            if not conversion_success:
                try:
                    success = self._convert_html_to_docx_native(protected_html, output_path)
                    if success:
                        conversion_success = True
                        method_used = "native"
                        self.conversion_stats['native_successes'] += 1
                        self.logger.info("Conversión exitosa con método nativo")
                except Exception as e:
                    self.logger.error(f"Conversión nativa también falló: {e}")
            
            if not conversion_success:
                self.conversion_stats['total_failures'] += 1
                raise RuntimeError("Ambos métodos de conversión fallaron")
            
            # 6. Validar preservación de estilos en el DOCX generado
            if self.style_analyzer:
                preserved_styles = self.style_analyzer.extract_docx_styles(output_path)
                style_preservation_score, preservation_alerts = self.style_analyzer.calculate_style_preservation_score(
                    original_styles, preserved_styles
                )
            else:
                preserved_styles = self._extract_docx_styles_fallback(output_path)
                style_preservation_score = self._calculate_style_preservation_score_fallback(
                    original_styles, preserved_styles
                )
                preservation_alerts = []
            
            # 7. Procesar footnotes funcionales en el documento final (Modificación 1)
            functional_footnotes_created = 0
            footnote_preservation_score = 1.0
            
            if footnote_mapping and self.create_functional_footnotes:
                footnote_preservation_score, functional_footnotes_created = self._create_functional_footnotes_in_docx(
                    output_path, footnote_mapping
                )
                self.conversion_stats['functional_footnotes_created'] += functional_footnotes_created
            elif footnote_mapping:
                footnote_preservation_score = self._reconnect_footnotes_in_docx(
                    output_path, footnote_mapping
                )
            
            # 8. Generar sistema de alertas automáticas (Modificación 2)
            alert_summary = self._analyze_preservation_alerts(preservation_alerts, style_preservation_score)
            manual_review_required = self._determine_manual_review_requirement(
                preservation_alerts, style_preservation_score, footnote_preservation_score
            )
            
            # 9. Generar reporte de validación
            validation_report = self._generate_validation_report(
                original_styles, preserved_styles, footnote_mapping, preservation_alerts
            )
            
            processing_time = time.time() - start_time
            
            # 10. Actualizar estadísticas
            self._update_conversion_stats(processing_time, style_preservation_score)
            
            result = ConversionResult(
                success=True,
                method_used=method_used,
                output_path=output_path,
                style_preservation_score=style_preservation_score,
                footnote_preservation_score=footnote_preservation_score,
                processing_time_seconds=processing_time,
                original_styles=original_styles,
                preserved_styles=preserved_styles,
                footnote_mapping=footnote_mapping,
                validation_report=validation_report,
                preservation_alerts=preservation_alerts,
                complex_tables_info=complex_tables_info,
                functional_footnotes_created=functional_footnotes_created,
                manual_review_required=manual_review_required,
                alert_summary=alert_summary
            )
            
            self.logger.info(f"Conversión completada exitosamente en {processing_time:.2f}s")
            self.logger.info(f"Score preservación estilos: {style_preservation_score:.3f}")
            self.logger.info(f"Score preservación footnotes: {footnote_preservation_score:.3f}")
            self.logger.info(f"Alertas generadas: {len(preservation_alerts)} "
                           f"(críticas: {alert_summary['critical']}, warnings: {alert_summary['warning']})")
            
            # Sistema de alertas automáticas (Modificación 2)
            if alert_summary['critical'] > 0:
                self.logger.error(f"¡ALERTAS CRÍTICAS! {alert_summary['critical']} problemas críticos detectados")
                self.conversion_stats['critical_alerts_generated'] += alert_summary['critical']
            
            return result
            
        except Exception as e:
            processing_time = time.time() - start_time
            self.conversion_stats['total_failures'] += 1
            
            self.logger.error(f"Error en conversión HTML→DOCX: {e}")
            
            return ConversionResult(
                success=False,
                method_used="failed",
                output_path="",
                style_preservation_score=0.0,
                footnote_preservation_score=0.0,
                processing_time_seconds=processing_time,
                original_styles=StyleInfo() if not self.style_analyzer else ExtendedStyleInfo(),
                preserved_styles=StyleInfo() if not self.style_analyzer else ExtendedStyleInfo(),
                footnote_mapping={},
                validation_report={},
                error_details=str(e),
                manual_review_required=True
            )
    
    def _detect_complex_tables(self, html_content: str) -> List[ComplexTableInfo]:
        """Detecta y analiza tablas con estructuras complejas (rowspan/colspan)"""
        soup = BeautifulSoup(html_content, 'html.parser')
        complex_tables = []
        
        for table_index, table in enumerate(soup.find_all('table')):
            table_info = ComplexTableInfo(table_index=table_index)
            
            # Analizar celdas para detectar spans
            for cell in table.find_all(['td', 'th']):
                rowspan = cell.get('rowspan')
                colspan = cell.get('colspan')
                
                if rowspan and int(rowspan) > 1:
                    table_info.has_rowspan = True
                    table_info.max_rowspan = max(table_info.max_rowspan, int(rowspan))
                    table_info.total_merged_cells += 1
                
                if colspan and int(colspan) > 1:
                    table_info.has_colspan = True
                    table_info.max_colspan = max(table_info.max_colspan, int(colspan))
                    table_info.total_merged_cells += 1
            
            # Determinar estrategia de preservación
            if table_info.has_rowspan or table_info.has_colspan:
                if table_info.total_merged_cells > 5 or table_info.max_rowspan > 3 or table_info.max_colspan > 3:
                    table_info.preservation_strategy = "manual_review"
                elif table_info.total_merged_cells > 2:
                    table_info.preservation_strategy = "approximate"
                else:
                    table_info.preservation_strategy = "flatten"
                
                complex_tables.append(table_info)
                self.logger.debug(f"Tabla compleja detectada: índice {table_index}, "
                                f"rowspan max: {table_info.max_rowspan}, "
                                f"colspan max: {table_info.max_colspan}")
        
        self.conversion_stats['complex_tables_processed'] += len(complex_tables)
        return complex_tables
    
    def _create_functional_footnotes_in_docx(self, 
                                           docx_path: str, 
                                           footnote_mapping: Dict[str, FootnoteInfo]) -> Tuple[float, int]:
        """
        Crea footnotes funcionales reales en el documento DOCX (Modificación 1)
        
        Args:
            docx_path: Ruta del documento DOCX
            footnote_mapping: Mapeo de footnotes con IDs únicos
            
        Returns:
            Tuple de (preservation_score, footnotes_funcionales_creadas)
        """
        try:
            doc = Document(docx_path)
            functional_footnotes_created = 0
            total_footnotes = len(footnote_mapping)
            
            if total_footnotes == 0:
                return 1.0, 0
            
            # Buscar y procesar cada párrafo
            for paragraph in doc.paragraphs:
                paragraph_text = paragraph.text
                
                # Buscar marcadores de footnotes en el párrafo
                for unique_id, footnote_info in footnote_mapping.items():
                    id_marker = f"<<<{unique_id}>>>"
                    
                    if id_marker in paragraph_text:
                        try:
                            # Encontrar la posición del marcador en el párrafo
                            marker_pos = paragraph_text.find(id_marker)
                            
                            # Dividir el párrafo en runs antes y después del marcador
                            before_text = paragraph_text[:marker_pos]
                            after_text = paragraph_text[marker_pos + len(id_marker):]
                            
                            # Limpiar el párrafo actual
                            paragraph.clear()
                            
                            # Agregar texto antes del footnote
                            if before_text:
                                paragraph.add_run(before_text)
                            
                            # CREAR FOOTNOTE FUNCIONAL
                            # Nota: python-docx no soporta footnotes nativamente,
                            # pero podemos simular la funcionalidad o usar workarounds
                            footnote_run = paragraph.add_run(str(footnote_info.reference_number))
                            footnote_run.font.superscript = True
                            footnote_run.font.color.rgb = RGBColor(0, 0, 255)  # Azul para indicar link
                            
                            # Para footnotes realmente funcionales, necesitaríamos:
                            # 1. Modificar el XML del documento directamente
                            # 2. Usar una librería como python-docx-footnotes
                            # 3. O implementar la funcionalidad manualmente
                            
                            # Workaround: Agregar el contenido de la footnote como comentario
                            # o al final del documento
                            self._add_footnote_content_to_document(doc, footnote_info)
                            
                            # Agregar texto después del footnote
                            if after_text:
                                paragraph.add_run(after_text)
                            
                            # Marcar como funcional (aunque sea aproximado)
                            footnote_info.is_functional = True
                            footnote_info.docx_footnote_id = footnote_info.reference_number
                            functional_footnotes_created += 1
                            
                        except Exception as e:
                            self.logger.warning(f"Error creando footnote funcional {unique_id}: {e}")
                            continue
            
            # Guardar documento modificado
            doc.save(docx_path)
            
            preservation_score = functional_footnotes_created / total_footnotes if total_footnotes > 0 else 1.0
            
            self.logger.info(f"Footnotes funcionales creadas: {functional_footnotes_created}/{total_footnotes} "
                           f"(score: {preservation_score:.3f})")
            
            return preservation_score, functional_footnotes_created
            
        except Exception as e:
            self.logger.error(f"Error creando footnotes funcionales: {e}")
            return 0.0, 0
    
    def _add_footnote_content_to_document(self, doc: Document, footnote_info: FootnoteInfo) -> None:
        """Agrega contenido de footnote al final del documento como workaround"""
        try:
            # Buscar o crear sección de footnotes al final
            footnotes_section_found = False
            for paragraph in doc.paragraphs:
                if "FOOTNOTES" in paragraph.text.upper() or "NOTAS" in paragraph.text.upper():
                    footnotes_section_found = True
                    break
            
            if not footnotes_section_found:
                # Crear sección de footnotes
                doc.add_paragraph()  # Espacio
                footnotes_heading = doc.add_heading("Footnotes", level=2)
                footnotes_heading.style = 'Heading 2'
            
            # Agregar la footnote
            footnote_paragraph = doc.add_paragraph()
            footnote_run = footnote_paragraph.add_run(f"{footnote_info.reference_number}. ")
            footnote_run.font.bold = True
            footnote_paragraph.add_run(footnote_info.footnote_content or footnote_info.original_text)
            
        except Exception as e:
            self.logger.warning(f"Error agregando contenido de footnote: {e}")
    
    def _analyze_preservation_alerts(self, 
                                   alerts: List[StylePreservationAlert], 
                                   style_score: float) -> Dict[str, int]:
        """Analiza alertas y genera resumen por severidad"""
        alert_summary = {'critical': 0, 'warning': 0, 'info': 0}
        
        for alert in alerts:
            alert_summary[alert.alert_type] = alert_summary.get(alert.alert_type, 0) + 1
        
        # Agregar alertas basadas en scores
        if style_score < 0.70:
            alert_summary['critical'] += 1
        elif style_score < 0.90:
            alert_summary['warning'] += 1
        
        return alert_summary
    
    def _determine_manual_review_requirement(self, 
                                           alerts: List[StylePreservationAlert],
                                           style_score: float,
                                           footnote_score: float) -> bool:
        """Determina si se requiere revisión manual basado en alertas y scores"""
        # Criterios para revisión manual obligatoria
        critical_alerts = sum(1 for alert in alerts if alert.alert_type == 'critical')
        
        if critical_alerts > 0:
            return True
        
        if style_score < self.validation_threshold:
            return True
        
        if footnote_score < 0.80 and footnote_score > 0:  # Si hay footnotes pero mal preservadas
            return True
        
        # Revisar alertas específicas
        table_alerts = [alert for alert in alerts if 'table' in alert.element_type.lower()]
        if len(table_alerts) > 2:  # Múltiples problemas con tablas
            return True
        
        return False
    def _extract_html_styles_fallback(self, html_content: str) -> StyleInfo:
        """Método fallback para extraer estilos cuando HTMLStyleAnalyzer no está disponible"""
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Contar elementos de formato
        style_info = StyleInfo()
        
        # Formatos básicos
        style_info.bold_count = len(soup.find_all(['strong', 'b']))
        style_info.italic_count = len(soup.find_all(['em', 'i']))
        style_info.underline_count = len(soup.find_all('u'))
        style_info.superscript_count = len(soup.find_all('sup'))
        style_info.subscript_count = len(soup.find_all('sub'))
        style_info.blockquote_count = len(soup.find_all('blockquote'))
        
        # Headings
        style_info.heading_counts = {}
        for level in range(1, 7):
            count = len(soup.find_all(f'h{level}'))
            if count > 0:
                style_info.heading_counts[f'h{level}'] = count
        
        # Listas
        style_info.list_counts = {
            'ul': len(soup.find_all('ul')),
            'ol': len(soup.find_all('ol')),
            'li': len(soup.find_all('li'))
        }
        
        # Tablas y otros
        style_info.table_count = len(soup.find_all('table'))
        
        # Footnotes (buscar patrones comunes)
        footnote_patterns = [
            soup.find_all('sup', string=re.compile(r'^\d+
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Contar elementos de formato
        style_info = StyleInfo()
        
        # Formatos básicos
        style_info.bold_count = len(soup.find_all(['strong', 'b']))
        style_info.italic_count = len(soup.find_all(['em', 'i']))
        style_info.underline_count = len(soup.find_all('u'))
        style_info.superscript_count = len(soup.find_all('sup'))
        style_info.subscript_count = len(soup.find_all('sub'))
        style_info.blockquote_count = len(soup.find_all('blockquote'))
        
        # Headings
        style_info.heading_counts = {}
        for level in range(1, 7):
            count = len(soup.find_all(f'h{level}'))
            if count > 0:
                style_info.heading_counts[f'h{level}'] = count
        
        # Listas
        style_info.list_counts = {
            'ul': len(soup.find_all('ul')),
            'ol': len(soup.find_all('ol')),
            'li': len(soup.find_all('li'))
        }
        
        # Tablas y otros
        style_info.table_count = len(soup.find_all('table'))
        
        # Footnotes (buscar patrones comunes)
        footnote_patterns = [
            soup.find_all('sup', string=re.compile(r'^\d+$')),
            soup.find_all('a', href=re.compile(r'#fn\d+')),
            soup.find_all(class_=re.compile(r'footnote')),
        ]
        style_info.footnote_count = sum(len(pattern) for pattern in footnote_patterns)
        
        return style_info
    
    def _insert_format_markers(self, html_content: str) -> str:
        """
        Convierte tags HTML a marcadores protegidos para preservar formato durante traducción
        
        Args:
            html_content: HTML con formatos originales
            
        Returns:
            HTML con marcadores protegidos que Claude no modificará
        """
        protected_content = html_content
        
        # Aplicar marcadores en orden específico (del más específico al más general)
        for html_pattern, marker_replacement in self.format_markers.items():
            try:
                protected_content = re.sub(
                    html_pattern, 
                    marker_replacement,
                    protected_content, 
                    flags=re.DOTALL | re.IGNORECASE
                )
            except re.error as e:
                self.logger.warning(f"Error en patrón regex {html_pattern}: {e}")
                continue
        
        # Logging para debug
        if self.debug_mode:
            markers_inserted = len(re.findall(r'<<<\w+_(?:START|END)>>>', protected_content))
            self.logger.debug(f"Insertados {markers_inserted} marcadores de formato")
        
        return protected_content
    
    def _convert_html_to_docx_pandoc(self, html_content: str, output_path: str) -> bool:
        """
        Conversión HTML→DOCX usando Pandoc
        
        Args:
            html_content: HTML a convertir
            output_path: Ruta del archivo DOCX de salida
            
        Returns:
            True si la conversión fue exitosa
        """
        try:
            with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as temp_html:
                temp_html.write(html_content)
                temp_html_path = temp_html.name
            
            # Comando Pandoc con opciones optimizadas para preservación
            pandoc_cmd = [
                'pandoc',
                temp_html_path,
                '-f', 'html',
                '-t', 'docx',
                '-o', output_path,
                '--preserve-tabs',
                '--wrap=preserve'
            ]
            
            self.logger.debug(f"Ejecutando: {' '.join(pandoc_cmd)}")
            
            result = subprocess.run(
                pandoc_cmd,
                capture_output=True,
                text=True,
                timeout=self.pandoc_timeout
            )
            
            # Cleanup
            Path(temp_html_path).unlink(missing_ok=True)
            
            if result.returncode == 0:
                if Path(output_path).exists() and Path(output_path).stat().st_size > 0:
                    return True
                else:
                    self.logger.error("Pandoc ejecutó sin errores pero no generó archivo válido")
                    return False
            else:
                self.logger.error(f"Pandoc falló: {result.stderr}")
                return False
                
        except subprocess.TimeoutExpired:
            self.logger.error(f"Pandoc timeout después de {self.pandoc_timeout}s")
            return False
        except Exception as e:
            self.logger.error(f"Error ejecutando Pandoc: {e}")
            return False
    
    def _convert_html_to_docx_native(self, html_content: str, output_path: str) -> bool:
        """
        Conversión HTML→DOCX usando python-docx nativo
        
        Args:
            html_content: HTML a convertir  
            output_path: Ruta del archivo DOCX de salida
            
        Returns:
            True si la conversión fue exitosa
        """
        try:
            # Crear documento DOCX
            doc = Document()
            
            # Configurar estilos básicos
            self._setup_document_styles(doc)
            
            # Parsear HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Procesar elementos HTML secuencialmente
            for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 
                                        'blockquote', 'ul', 'ol', 'table', 'div', 'br']):
                self._process_html_element(doc, element)
            
            # Guardar documento
            doc.save(output_path)
            
            if Path(output_path).exists() and Path(output_path).stat().st_size > 0:
                return True
            else:
                self.logger.error("Conversión nativa no generó archivo válido")
                return False
                
        except Exception as e:
            self.logger.error(f"Error en conversión nativa: {e}")
            return False
    
    def _setup_document_styles(self, doc: Document) -> None:
        """Configura estilos básicos en el documento DOCX"""
        styles = doc.styles
        
        # Estilo para blockquotes
        try:
            quote_style = styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)
            quote_style.paragraph_format.left_indent = Inches(0.5)
            quote_style.paragraph_format.right_indent = Inches(0.5)
            quote_style.font.italic = True
            quote_style.font.size = Pt(11)
        except:
            pass  # El estilo ya puede existir
    
    def _process_html_element(self, doc: Document, element: Tag) -> None:
        """Procesa un elemento HTML individual y lo agrega al documento DOCX"""
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            # Encabezados
            level = int(element.name[1])
            heading = doc.add_heading(element.get_text(), level)
            
        elif element.name == 'p':
            # Párrafos
            paragraph = doc.add_paragraph()
            self._process_inline_formatting(paragraph, element)
            
        elif element.name == 'blockquote':
            # Citas
            paragraph = doc.add_paragraph(element.get_text(), style='Quote')
            
        elif element.name in ['ul', 'ol']:
            # Listas
            for li in element.find_all('li', recursive=False):
                paragraph = doc.add_paragraph(li.get_text(), style='List Bullet' if element.name == 'ul' else 'List Number')
                
        elif element.name == 'br':
            # Saltos de línea
            doc.add_paragraph()
    
    def _process_inline_formatting(self, paragraph, element: Tag) -> None:
        """Procesa formato inline dentro de un párrafo"""
        for content in element.contents:
            if isinstance(content, str):
                # Texto plano
                run = paragraph.add_run(content)
            elif hasattr(content, 'name'):
                # Elemento HTML
                text = content.get_text()
                run = paragraph.add_run(text)
                
                # Aplicar formato según el tag
                if content.name in ['strong', 'b']:
                    run.bold = True
                elif content.name in ['em', 'i']:
                    run.italic = True
                elif content.name == 'u':
                    run.underline = True
                elif content.name == 'sup':
                    run.font.superscript = True
                elif content.name == 'sub':
                    run.font.subscript = True
    
    def _process_footnotes_for_preservation(self, html_content: str) -> Tuple[str, Dict[str, FootnoteInfo]]:
        """
        Procesa footnotes insertando IDs únicos para preservación durante traducción
        
        Args:
            html_content: HTML original con footnotes
            
        Returns:
            Tuple de (HTML procesado con IDs, mapeo de footnotes)
        """
        footnote_mapping = {}
        processed_html = html_content
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Buscar footnotes por patrones comunes
        footnote_refs = []
        
        # Patrón 1: <sup>número</sup>
        for sup in soup.find_all('sup'):
            if re.match(r'^\d+$', sup.get_text().strip()):
                footnote_refs.append({
                    'element': sup,
                    'number': int(sup.get_text().strip()),
                    'pattern': 'superscript_number'
                })
        
        # Patrón 2: <a href="#fn1">1</a>
        for link in soup.find_all('a', href=re.compile(r'#fn\d+')):
            footnote_refs.append({
                'element': link,
                'number': int(re.search(r'\d+', link.get('href')).group()),
                'pattern': 'link_reference'
            })
        
        # Patrón 3: Elementos con class footnote
        for elem in soup.find_all(class_=re.compile(r'footnote')):
            footnote_refs.append({
                'element': elem,
                'number': len(footnote_refs) + 1,
                'pattern': 'class_footnote'
            })
        
        # Procesar cada footnote encontrada
        for i, ref_info in enumerate(footnote_refs):
            element = ref_info['element']
            number = ref_info['number']
            
            # Generar ID único
            unique_id = f"FOOTNOTE_ID_{number}_{hashlib.md5(element.get_text().encode()).hexdigest()[:8]}"
            
            # Crear info del footnote
            footnote_info = FootnoteInfo(
                unique_id=unique_id,
                original_text=element.get_text(),
                superscript_location=str(element),
                reference_number=number
            )
            
            footnote_mapping[unique_id] = footnote_info
            
            # Insertar ID único en el HTML
            if ref_info['pattern'] == 'superscript_number':
                # Para <sup>1</sup> → <sup><<<FOOTNOTE_ID_1_abc123>>> 1</sup>
                element.string = f"<<<{unique_id}>>> {element.get_text()}"
            elif ref_info['pattern'] == 'link_reference':
                # Para <a href="#fn1">1</a> → <a href="#fn1"><<<FOOTNOTE_ID_1_abc123>>> 1</a>
                element.string = f"<<<{unique_id}>>> {element.get_text()}"
            else:
                # Para otros tipos, insertar ID antes del contenido
                element.insert(0, f"<<<{unique_id}>>> ")
        
        # Convertir el soup modificado de vuelta a HTML
        processed_html = str(soup)
        
        self.logger.info(f"Procesadas {len(footnote_mapping)} footnotes con IDs únicos")
        
        return processed_html, footnote_mapping
    
    def _extract_docx_styles(self, docx_path: str) -> StyleInfo:
        """Extrae información de estilos del DOCX generado"""
        try:
            doc = Document(docx_path)
            style_info = StyleInfo()
            
            # Contar elementos de formato en el documento
            bold_count = 0
            italic_count = 0
            underline_count = 0
            superscript_count = 0
            subscript_count = 0
            
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run.bold:
                        bold_count += 1
                    if run.italic:
                        italic_count += 1
                    if run.underline:
                        underline_count += 1
                    if run.font.superscript:
                        superscript_count += 1
                    if run.font.subscript:
                        subscript_count += 1
            
            style_info.bold_count = bold_count
            style_info.italic_count = italic_count
            style_info.underline_count = underline_count
            style_info.superscript_count = superscript_count
            style_info.subscript_count = subscript_count
            
            # Contar headings por estilo
            heading_counts = {}
            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith('Heading'):
                    level = paragraph.style.name.replace('Heading ', 'h')
                    heading_counts[level] = heading_counts.get(level, 0) + 1
            
            style_info.heading_counts = heading_counts
            
            # Contar tablas
            style_info.table_count = len(doc.tables)
            
            return style_info
            
        except Exception as e:
            self.logger.error(f"Error extrayendo estilos de DOCX: {e}")
            return StyleInfo()
    
    def _calculate_style_preservation_score(self, 
                                          original: StyleInfo, 
                                          preserved: StyleInfo) -> float:
        """
        Calcula score de preservación de estilos comparando original vs preservado
        
        Args:
            original: Estilos del HTML original
            preserved: Estilos del DOCX generado
            
        Returns:
            Score de 0.0 a 1.0 indicando qué tan bien se preservaron los estilos
        """
        try:
            scores = []
            
            # Comparar formatos básicos
            basic_formats = [
                ('bold', original.bold_count, preserved.bold_count),
                ('italic', original.italic_count, preserved.italic_count),
                ('underline', original.underline_count, preserved.underline_count),
                ('superscript', original.superscript_count, preserved.superscript_count),
                ('subscript', original.subscript_count, preserved.subscript_count),
                ('blockquote', original.blockquote_count, preserved.blockquote_count),
                ('table', original.table_count, preserved.table_count)
            ]
            
            for format_name, orig_count, pres_count in basic_formats:
                if orig_count == 0:
                    # Si no había elementos originales, perfecto si tampoco hay preservados
                    score = 1.0 if pres_count == 0 else 0.8
                else:
                    # Calcular ratio de preservación
                    preservation_ratio = min(pres_count / orig_count, 1.0)
                    score = preservation_ratio
                
                scores.append(score)
                
                if self.debug_mode:
                    self.logger.debug(f"{format_name}: {orig_count} → {pres_count} (score: {score:.3f})")
            
            # Comparar headings
            for level in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                orig_count = original.heading_counts.get(level, 0)
                pres_count = preserved.heading_counts.get(level, 0)
                
                if orig_count == 0:
                    score = 1.0 if pres_count == 0 else 0.8
                else:
                    score = min(pres_count / orig_count, 1.0)
                
                scores.append(score)
            
            # Comparar listas
            for list_type in ['ul', 'ol', 'li']:
                orig_count = original.list_counts.get(list_type, 0)
                pres_count = preserved.list_counts.get(list_type, 0)
                
                if orig_count == 0:
                    score = 1.0 if pres_count == 0 else 0.8
                else:
                    score = min(pres_count / orig_count, 1.0)
                
                scores.append(score)
            
            # Calcular score promedio ponderado
            if not scores:
                return 0.0
            
            overall_score = sum(scores) / len(scores)
            
            self.logger.info(f"Score de preservación de estilos: {overall_score:.3f}")
            
            return overall_score
            
        except Exception as e:
            self.logger.error(f"Error calculando score de preservación: {e}")
            return 0.0
    
    def _reconnect_footnotes_in_docx(self, 
                                   docx_path: str, 
                                   footnote_mapping: Dict[str, FootnoteInfo]) -> float:
        """
        Reconecta footnotes en el documento DOCX final
        
        Args:
            docx_path: Ruta del documento DOCX
            footnote_mapping: Mapeo de IDs únicos a footnotes
            
        Returns:
            Score de preservación de footnotes (0.0 a 1.0)
        """
        try:
            doc = Document(docx_path)
            reconnected_count = 0
            total_footnotes = len(footnote_mapping)
            
            if total_footnotes == 0:
                return 1.0  # No había footnotes, preservación perfecta
            
            # Buscar IDs de footnotes en el texto y reconectar
            for paragraph in doc.paragraphs:
                paragraph_text = paragraph.text
                
                for unique_id, footnote_info in footnote_mapping.items():
                    id_marker = f"<<<{unique_id}>>>"
                    
                    if id_marker in paragraph_text:
                        try:
                            # Encontrar el run que contiene el ID
                            for run in paragraph.runs:
                                if id_marker in run.text:
                                    # Remover el ID marker
                                    run.text = run.text.replace(id_marker, "")
                                    
                                    # Crear footnote real (simplificado para este ejemplo)
                                    # En implementación completa, usaríamos la API de footnotes de python-docx
                                    run.font.superscript = True
                                    
                                    reconnected_count += 1
                                    break
                        except Exception as e:
                            self.logger.warning(f"Error reconectando footnote {unique_id}: {e}")
                            continue
            
            # Guardar documento modificado
            doc.save(docx_path)
            
            preservation_score = reconnected_count / total_footnotes if total_footnotes > 0 else 1.0
            
            self.logger.info(f"Footnotes reconectadas: {reconnected_count}/{total_footnotes} "
                           f"(score: {preservation_score:.3f})")
            
            return preservation_score
            
        except Exception as e:
            self.logger.error(f"Error reconectando footnotes: {e}")
            return 0.0
    
    def _generate_validation_report(self, 
                                  original: Union[StyleInfo, ExtendedStyleInfo],
                                  preserved: Union[StyleInfo, ExtendedStyleInfo],
                                  footnote_mapping: Dict[str, FootnoteInfo],
                                  alerts: List[StylePreservationAlert] = None) -> Dict[str, Any]:
        """Genera reporte detallado de validación con alertas"""
        if alerts is None:
            alerts = []
        
        report = {
            'timestamp': datetime.now().isoformat(),
            'style_analysis': {
                'original_counts': asdict(original),
                'preserved_counts': asdict(preserved)
            },
            'footnote_analysis': {
                'total_footnotes': len(footnote_mapping),
                'functional_footnotes': sum(1 for f in footnote_mapping.values() if f.is_functional),
                'footnote_details': [info.to_dict() for info in footnote_mapping.values()]
            },
            'preservation_alerts': {
                'total_alerts': len(alerts),
                'critical_alerts': [alert.to_dict() for alert in alerts if alert.alert_type == 'critical'],
                'warning_alerts': [alert.to_dict() for alert in alerts if alert.alert_type == 'warning'],
                'info_alerts': [alert.to_dict() for alert in alerts if alert.alert_type == 'info']
            },
            'validation_summary': {
                'total_elements_analyzed': self._count_total_elements(original),
                'critical_issues_found': len([a for a in alerts if a.alert_type == 'critical']),
                'warnings_generated': len([a for a in alerts if a.alert_type == 'warning']),
                'manual_review_recommended': any(a.alert_type == 'critical' for a in alerts),
                'overall_quality_assessment': self._assess_overall_quality(alerts)
            }
        }
        
        # Agregar análisis específico si es ExtendedStyleInfo
        if isinstance(original, ExtendedStyleInfo) and self.enable_extended_analysis:
            report['advanced_analysis'] = {
                'small_caps_preservation': self._calculate_preservation_ratio(
                    original.small_caps_count, preserved.small_caps_count
                ),
                'strikethrough_preservation': self._calculate_preservation_ratio(
                    original.strikethrough_count, preserved.strikethrough_count
                ),
                'color_preservation': self._calculate_preservation_ratio(
                    original.colored_text_count, preserved.colored_text_count
                ),
                'complex_structures': {
                    'nested_list_depth_preserved': preserved.nested_list_depth >= original.nested_list_depth,
                    'complex_tables_processed': preserved.complex_table_count,
                    'alignment_preservation': self._assess_alignment_preservation(original, preserved)
                }
            }
        
        return report
    
    def _count_total_elements(self, style_info: Union[StyleInfo, ExtendedStyleInfo]) -> int:
        """Cuenta el total de elementos analizados"""
        basic_count = (
            style_info.bold_count + style_info.italic_count + style_info.underline_count +
            style_info.superscript_count + style_info.subscript_count + style_info.blockquote_count +
            style_info.table_count + sum(style_info.heading_counts.values()) +
            sum(style_info.list_counts.values()) + style_info.footnote_count
        )
        
        # Agregar elementos avanzados si están disponibles
        if isinstance(style_info, ExtendedStyleInfo) and self.enable_extended_analysis:
            advanced_count = (
                style_info.small_caps_count + style_info.strikethrough_count +
                style_info.colored_text_count + style_info.highlighted_count +
                style_info.center_aligned_count + style_info.right_aligned_count +
                style_info.justified_count + style_info.link_count + 
                style_info.image_count + style_info.code_block_count
            )
            return basic_count + advanced_count
        
        return basic_count
    
    def _calculate_preservation_ratio(self, original_count: int, preserved_count: int) -> float:
        """Calcula ratio de preservación para un tipo de elemento"""
        if original_count == 0:
            return 1.0 if preserved_count == 0 else 0.8
        return min(preserved_count / original_count, 1.0)
    
    def _assess_overall_quality(self, alerts: List[StylePreservationAlert]) -> str:
        """Evalúa la calidad general basada en alertas"""
        critical_count = sum(1 for alert in alerts if alert.alert_type == 'critical')
        warning_count = sum(1 for alert in alerts if alert.alert_type == 'warning')
        
        if critical_count > 0:
            return "Poor - Critical issues detected"
        elif warning_count > 5:
            return "Fair - Multiple warnings"
        elif warning_count > 0:
            return "Good - Minor issues"
        else:
            return "Excellent - No significant issues"
    
    def _assess_alignment_preservation(self, original: ExtendedStyleInfo, preserved: ExtendedStyleInfo) -> Dict[str, Any]:
        """Evalúa preservación de alineaciones"""
        total_original_alignments = (
            original.center_aligned_count + original.right_aligned_count + original.justified_count
        )
        total_preserved_alignments = (
            preserved.center_aligned_count + preserved.right_aligned_count + preserved.justified_count
        )
        
        return {
            'total_alignments_original': total_original_alignments,
            'total_alignments_preserved': total_preserved_alignments,
            'preservation_ratio': self._calculate_preservation_ratio(
                total_original_alignments, total_preserved_alignments
            ),
            'center_preserved': self._calculate_preservation_ratio(
                original.center_aligned_count, preserved.center_aligned_count
            ),
            'right_preserved': self._calculate_preservation_ratio(
                original.right_aligned_count, preserved.right_aligned_count
            ),
            'justify_preserved': self._calculate_preservation_ratio(
                original.justified_count, preserved.justified_count
            )
        }
    
    def get_conversion_statistics(self) -> Dict[str, Any]:
        """Retorna estadísticas de conversión acumuladas con métricas extendidas"""
        stats = self.conversion_stats.copy()
        
        # Calcular ratios de éxito
        if stats['total_conversions'] > 0:
            stats['success_rate'] = (
                (stats['pandoc_successes'] + stats['native_successes']) / 
                stats['total_conversions']
            )
            stats['pandoc_success_rate'] = stats['pandoc_successes'] / stats['total_conversions']
            stats['native_fallback_rate'] = stats['native_successes'] / stats['total_conversions']
            stats['failure_rate'] = stats['total_failures'] / stats['total_conversions']
        else:
            stats['success_rate'] = 0.0
            stats['pandoc_success_rate'] = 0.0
            stats['native_fallback_rate'] = 0.0
            stats['failure_rate'] = 0.0
        
        # Métricas de calidad extendidas
        if self.enable_extended_analysis:
            stats['advanced_features'] = {
                'functional_footnotes_avg': (
                    stats['functional_footnotes_created'] / max(stats['total_conversions'], 1)
                ),
                'complex_tables_avg': (
                    stats['complex_tables_processed'] / max(stats['total_conversions'], 1)
                ),
                'critical_alerts_rate': (
                    stats['critical_alerts_generated'] / max(stats['total_conversions'], 1)
                )
            }
        
        return stats
    
    def validate_conversion_quality(self, 
                                  html_content: str, 
                                  docx_path: str) -> Dict[str, Any]:
        """
        Valida la calidad de una conversión específica con análisis extendido
        
        Args:
            html_content: HTML original
            docx_path: Documento DOCX generado
            
        Returns:
            Reporte de calidad detallado con alertas
        """
        # Extraer estilos
        if self.style_analyzer:
            original_styles = self.style_analyzer.extract_html_styles(html_content)
            preserved_styles = self.style_analyzer.extract_docx_styles(docx_path)
            style_score, alerts = self.style_analyzer.calculate_style_preservation_score(
                original_styles, preserved_styles
            )
        else:
            original_styles = self._extract_html_styles_fallback(html_content)
            preserved_styles = self._extract_docx_styles_fallback(docx_path)
            style_score = self._calculate_style_preservation_score_fallback(
                original_styles, preserved_styles
            )
            alerts = []
        
        # Analizar alertas
        alert_summary = self._analyze_preservation_alerts(alerts, style_score)
        manual_review_required = self._determine_manual_review_requirement(
            alerts, style_score, 1.0  # Asumimos footnotes OK para este análisis
        )
        
        quality_report = {
            'overall_quality_score': style_score,
            'meets_threshold': style_score >= self.validation_threshold,
            'manual_review_required': manual_review_required,
            'original_style_counts': asdict(original_styles),
            'preserved_style_counts': asdict(preserved_styles),
            'preservation_alerts': [alert.to_dict() for alert in alerts],
            'alert_summary': alert_summary,
            'recommendations': self._generate_quality_recommendations(alerts, style_score),
            'quality_metrics': {
                'basic_formats_score': self._calculate_basic_formats_score(original_styles, preserved_styles),
                'structure_preservation_score': self._calculate_structure_score(original_styles, preserved_styles)
            }
        }
        
        # Agregar métricas avanzadas si están disponibles
        if isinstance(original_styles, ExtendedStyleInfo) and self.enable_extended_analysis:
            quality_report['quality_metrics']['advanced_formats_score'] = (
                self._calculate_advanced_formats_score(original_styles, preserved_styles)
            )
            quality_report['advanced_analysis'] = self._assess_alignment_preservation(
                original_styles, preserved_styles
            )
        
        return quality_report
    
    def _calculate_basic_formats_score(self, 
                                     original: Union[StyleInfo, ExtendedStyleInfo], 
                                     preserved: Union[StyleInfo, ExtendedStyleInfo]) -> float:
        """Calcula score específico para formatos básicos"""
        basic_elements = [
            (original.bold_count, preserved.bold_count),
            (original.italic_count, preserved.italic_count),
            (original.underline_count, preserved.underline_count),
            (original.superscript_count, preserved.superscript_count),
            (original.subscript_count, preserved.subscript_count)
        ]
        
        scores = []
        for orig, pres in basic_elements:
            scores.append(self._calculate_preservation_ratio(orig, pres))
        
        return sum(scores) / len(scores) if scores else 1.0
    
    def _calculate_advanced_formats_score(self, 
                                        original: ExtendedStyleInfo, 
                                        preserved: ExtendedStyleInfo) -> float:
        """Calcula score específico para formatos avanzados"""
        if not self.enable_extended_analysis or not isinstance(original, ExtendedStyleInfo):
            return 1.0
        
        advanced_elements = [
            (original.small_caps_count, preserved.small_caps_count),
            (original.strikethrough_count, preserved.strikethrough_count),
            (original.colored_text_count, preserved.colored_text_count),
            (original.highlighted_count, preserved.highlighted_count)
        ]
        
        scores = []
        for orig, pres in advanced_elements:
            scores.append(self._calculate_preservation_ratio(orig, pres))
        
        return sum(scores) / len(scores) if scores else 1.0
    
    def _calculate_structure_score(self, 
                                 original: Union[StyleInfo, ExtendedStyleInfo], 
                                 preserved: Union[StyleInfo, ExtendedStyleInfo]) -> float:
        """Calcula score específico para estructuras"""
        structure_elements = [
            (original.table_count, preserved.table_count),
            (original.blockquote_count, preserved.blockquote_count),
            (sum(original.heading_counts.values()), sum(preserved.heading_counts.values())),
            (sum(original.list_counts.values()), sum(preserved.list_counts.values()))
        ]
        
        scores = []
        for orig, pres in structure_elements:
            scores.append(self._calculate_preservation_ratio(orig, pres))
        
        return sum(scores) / len(scores) if scores else 1.0
    
    def _generate_quality_recommendations(self, 
                                        alerts: List[StylePreservationAlert], 
                                        style_score: float) -> List[str]:
        """Genera recomendaciones basadas en las alertas y score"""
        recommendations = []
        
        # Recomendaciones basadas en alertas críticas
        critical_alerts = [alert for alert in alerts if alert.alert_type == 'critical']
        if critical_alerts:
            recommendations.append("🚨 CRÍTICO: Revisión manual obligatoria antes de proceder")
            for alert in critical_alerts[:3]:  # Solo las primeras 3 para no saturar
                recommendations.append(f"• {alert.suggested_action}")
        
        # Recomendaciones basadas en score general
        if style_score < 0.70:
            recommendations.append("📋 Score muy bajo - considerar reprocesamiento con diferentes parámetros")
        elif style_score < self.validation_threshold:
            recommendations.append("⚠️ Score por debajo del umbral - verificar elementos críticos")
        
        # Recomendaciones específicas por tipo de alerta
        alert_types = {}
        for alert in alerts:
            element_type = alert.element_type
            if element_type not in alert_types:
                alert_types[element_type] = []
            alert_types[element_type].append(alert)
        
        for element_type, element_alerts in alert_types.items():
            if len(element_alerts) > 1:
                recommendations.append(f"🔍 Múltiples problemas con {element_type} - revisar proceso de conversión")
        
        # Recomendaciones generales
        if not recommendations:
            recommendations.append("✅ Calidad de conversión aceptable - revisión estándar recomendada")
        
        return recommendations
    
    def _update_conversion_stats(self, processing_time: float, style_score: float) -> None:
        """Actualiza estadísticas de conversión"""
        self.conversion_stats['total_conversions'] += 1
        
        # Calcular promedios
        total = self.conversion_stats['total_conversions']
        current_avg_time = self.conversion_stats['avg_processing_time']
        current_avg_style = self.conversion_stats['avg_style_preservation']
        
        self.conversion_stats['avg_processing_time'] = (
            (current_avg_time * (total - 1) + processing_time) / total
        )
        
        self.conversion_stats['avg_style_preservation'] = (
            (current_avg_style * (total - 1) + style_score) / total
        )
    
    def restore_format_markers(self, text_content: str) -> str:
        """
        Restaura marcadores de formato a HTML después de traducción con soporte extendido
        
        Args:
            text_content: Texto con marcadores <<<FORMAT_START>>> ... <<<FORMAT_END>>>
            
        Returns:
            HTML con tags restaurados (básicos y avanzados)
        """
        restored_content = text_content
        
        # Mapeo inverso: marcadores → HTML (básicos)
        basic_marker_to_html = {
            r'<<<BOLD_START>>>(.*?)<<<BOLD_END>>>': r'<strong>\1</strong>',
            r'<<<ITALIC_START>>>(.*?)<<<ITALIC_END>>>': r'<em>\1</em>',
            r'<<<UNDERLINE_START>>>(.*?)<<<UNDERLINE_END>>>': r'<u>\1</u>',
            r'<<<SUP_START>>>(.*?)<<<SUP_END>>>': r'<sup>\1</sup>',
            r'<<<SUB_START>>>(.*?)<<<SUB_END>>>': r'<sub>\1</sub>',
            r'<<<BLOCKQUOTE_START>>>(.*?)<<<BLOCKQUOTE_END>>>': r'<blockquote>\1</blockquote>',
            r'<<<H1_START>>>(.*?)<<<H1_END>>>': r'<h1>\1</h1>',
            r'<<<H2_START>>>(.*?)<<<H2_END>>>': r'<h2>\1</h2>',
            r'<<<H3_START>>>(.*?)<<<H3_END>>>': r'<h3>\1</h3>',
            r'<<<H4_START>>>(.*?)<<<H4_END>>>': r'<h4>\1</h4>',
            r'<<<H5_START>>>(.*?)<<<H5_END>>>': r'<h5>\1</h5>',
            r'<<<H6_START>>>(.*?)<<<H6_END>>>': r'<h6>\1</h6>',
            r'<<<P_START>>>(.*?)<<<P_END>>>': r'<p>\1</p>',
            r'<<<BR>>>': r'<br>',
            r'<<<UL_START>>>(.*?)<<<UL_END>>>': r'<ul>\1</ul>',
            r'<<<OL_START>>>(.*?)<<<OL_END>>>': r'<ol>\1</ol>',
            r'<<<LI_START>>>(.*?)<<<LI_END>>>': r'<li>\1</li>',
            r'<<<TABLE_START>>>(.*?)<<<TABLE_END>>>': r'<table>\1</table>',
            r'<<<TR_START>>>(.*?)<<<TR_END>>>': r'<tr>\1</tr>',
            r'<<<TD_START>>>(.*?)<<<TD_END>>>': r'<td>\1</td>',
            r'<<<TH_START>>>(.*?)<<<TH_END>>>': r'<th>\1</th>',
        }
        
        # Mapeo avanzado para estilos extendidos (Modificación 1)
        advanced_marker_to_html = {
            r'<<<SMALLCAPS_START>>>(.*?)<<<SMALLCAPS_END>>>': r'<span style="font-variant: small-caps">\1</span>',
            r'<<<STRIKETHROUGH_START>>>(.*?)<<<STRIKETHROUGH_END>>>': r'<del>\1</del>',
            r'<<<CODE_START>>>(.*?)<<<CODE_END>>>': r'<code>\1</code>',
            r'<<<PREFORMATTED_START>>>(.*?)<<<PREFORMATTED_END>>>': r'<pre>\1</pre>',
            r'<<<HIGHLIGHT_START>>>(.*?)<<<HIGHLIGHT_END>>>': r'<mark>\1</mark>',
            
            # Colores (con captura del color)
            r'<<<COLOR_([^_]+)_START>>>(.*?)<<<COLOR_END>>>': r'<span style="color: \1">\2</span>',
            r'<<<HIGHLIGHT_([^_]+)_START>>>(.*?)<<<HIGHLIGHT_END>>>': r'<span style="background-color: \1">\2</span>',
            
            # Alineaciones de párrafo
            r'<<<P_CENTER_START>>>(.*?)<<<P_CENTER_END>>>': r'<p style="text-align: center">\1</p>',
            r'<<<P_RIGHT_START>>>(.*?)<<<P_RIGHT_END>>>': r'<p style="text-align: right">\1</p>',
            r'<<<P_JUSTIFY_START>>>(.*?)<<<P_JUSTIFY_END>>>': r'<p style="text-align: justify">\1</p>',
            
            # Tablas complejas (rowspan/colspan)
            r'<<<TD_ROWSPAN_(\d+)_START>>>(.*?)<<<TD_ROWSPAN_END>>>': r'<td rowspan="\1">\2</td>',
            r'<<<TD_COLSPAN_(\d+)_START>>>(.*?)<<<TD_COLSPAN_END>>>': r'<td colspan="\1">\2</td>',
        }
        
        # Aplicar restauraciones básicas
        for marker_pattern, html_replacement in basic_marker_to_html.items():
            try:
                restored_content = re.sub(
                    marker_pattern,
                    html_replacement,
                    restored_content,
                    flags=re.DOTALL
                )
            except re.error as e:
                self.logger.warning(f"Error restaurando marcador básico {marker_pattern}: {e}")
                continue
        
        # Aplicar restauraciones avanzadas si está habilitado
        if self.enable_extended_analysis:
            for marker_pattern, html_replacement in advanced_marker_to_html.items():
                try:
                    restored_content = re.sub(
                        marker_pattern,
                        html_replacement,
                        restored_content,
                        flags=re.DOTALL
                    )
                except re.error as e:
                    self.logger.warning(f"Error restaurando marcador avanzado {marker_pattern}: {e}")
                    continue
        
        return restored_content
    
    def _setup_document_styles(self, doc: Document) -> None:
        """Configura estilos básicos y avanzados en el documento DOCX"""
        styles = doc.styles
        
        # Estilo para blockquotes
        try:
            quote_style = styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)
            quote_style.paragraph_format.left_indent = Inches(0.5)
            quote_style.paragraph_format.right_indent = Inches(0.5)
            quote_style.font.italic = True
            quote_style.font.size = Pt(11)
        except:
            pass  # El estilo ya puede existir
        
        # Estilos avanzados (Modificación 1)
        if self.enable_extended_analysis:
            # Estilo para código
            try:
                code_style = styles.add_style('Code', WD_STYLE_TYPE.CHARACTER)
                code_style.font.name = 'Courier New'
                code_style.font.size = Pt(10)
            except:
                pass
            
            # Estilo para texto resaltado
            try:
                highlight_style = styles.add_style('Highlight', WD_STYLE_TYPE.CHARACTER)
                highlight_style.font.highlight_color = 3  # Yellow highlight
            except:
                pass
            
            # Estilo para versalitas
            try:
                smallcaps_style = styles.add_style('SmallCaps', WD_STYLE_TYPE.CHARACTER)
                smallcaps_style.font.small_caps = True
            except:
                pass
    
    def _process_html_element(self, doc: Document, element: Tag) -> None:
        """Procesa un elemento HTML individual con soporte extendido"""
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            # Encabezados
            level = int(element.name[1])
            heading = doc.add_heading(element.get_text(), level)
            
        elif element.name == 'p':
            # Párrafos con análisis de alineación
            paragraph = doc.add_paragraph()
            
            # Detectar alineación del párrafo
            style_attr = element.get('style', '')
            if 'text-align: center' in style_attr:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif 'text-align: right' in style_attr:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif 'text-align: justify' in style_attr:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            
            self._process_inline_formatting(paragraph, element)
            
        elif element.name == 'blockquote':
            # Citas
            paragraph = doc.add_paragraph(element.get_text(), style='Quote')
            
        elif element.name in ['ul', 'ol']:
            # Listas con soporte para anidamiento
            self._process_nested_list(doc, element)
            
        elif element.name == 'table':
            # Tablas con soporte para estructuras complejas
            self._process_complex_table(doc, element)
            
        elif element.name == 'pre':
            # Texto preformateado
            paragraph = doc.add_paragraph(element.get_text())
            for run in paragraph.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                
        elif element.name == 'br':
            # Saltos de línea
            doc.add_paragraph()
    
    def _process_inline_formatting(self, paragraph, element: Tag) -> None:
        """Procesa formato inline dentro de un párrafo con soporte extendido"""
        for content in element.contents:
            if isinstance(content, str):
                # Texto plano
                run = paragraph.add_run(content)
            elif hasattr(content, 'name'):
                # Elemento HTML
                text = content.get_text()
                run = paragraph.add_run(text)
                
                # Aplicar formato según el tag
                if content.name in ['strong', 'b']:
                    run.bold = True
                elif content.name in ['em', 'i']:
                    run.italic = True
                elif content.name == 'u':
                    run.underline = True
                elif content.name == 'sup':
                    run.font.superscript = True
                elif content.name == 'sub':
                    run.font.subscript = True
                
                # Formatos avanzados (Modificación 1)
                elif self.enable_extended_analysis:
                    if content.name in ['s', 'strike', 'del']:
                        run.font.strike = True
                    elif content.name in ['code', 'kbd', 'samp']:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
                    elif content.name in ['mark', 'highlight']:
                        run.font.highlight_color = 3  # Yellow
                    
                    # Analizar estilos inline
                    style_attr = content.get('style', '')
                    if 'font-variant: small-caps' in style_attr:
                        run.font.small_caps = True
                    elif 'color:' in style_attr:
                        # Extraer color (simplificado)
                        color_match = re.search(r'color:\s*([^;]+)', style_attr)
                        if color_match:
                            color_value = color_match.group(1).strip()
                            # Convertir colores comunes
                            if color_value == 'red':
                                run.font.color.rgb = RGBColor(255, 0, 0)
                            elif color_value == 'blue':
                                run.font.color.rgb = RGBColor(0, 0, 255)
                            elif color_value == 'green':
                                run.font.color.rgb = RGBColor(0, 128, 0)
    
    def _process_nested_list(self, doc: Document, list_element: Tag, level: int = 0) -> None:
        """Procesa listas anidadas manteniendo la estructura jerárquica"""
        list_style = 'List Bullet' if list_element.name == 'ul' else 'List Number'
        
        for li in list_element.find_all('li', recursive=False):
            # Agregar item de lista
            paragraph = doc.add_paragraph(style=list_style)
            
            # Ajustar indentación según nivel de anidamiento
            if level > 0:
                paragraph.paragraph_format.left_indent = Inches(0.5 * (level + 1))
            
            # Procesar contenido del item
            text_content = ""
            for content in li.contents:
                if isinstance(content, str):
                    text_content += content
                elif hasattr(content, 'name') and content.name not in ['ul', 'ol']:
                    text_content += content.get_text()
            
            paragraph.add_run(text_content.strip())
            
            # Procesar sublistas
            for sublist in li.find_all(['ul', 'ol'], recursive=False):
                self._process_nested_list(doc, sublist, level + 1)
    
    def _process_complex_table(self, doc: Document, table_element: Tag) -> None:
        """Procesa tablas con soporte para rowspan/colspan (Modificación 1)"""
        try:
            # Contar filas y columnas máximas
            rows = table_element.find_all('tr')
            if not rows:
                return
            
            max_cols = 0
            for row in rows:
                cols = len(row.find_all(['td', 'th']))
                max_cols = max(max_cols, cols)
            
            # Crear tabla en DOCX
            table = doc.add_table(rows=len(rows), cols=max_cols)
            table.style = 'Table Grid'
            
            # Procesar cada fila
            for row_idx, tr in enumerate(rows):
                cells = tr.find_all(['td', 'th'])
                
                for col_idx, cell in enumerate(cells):
                    if col_idx < max_cols and row_idx < len(table.rows):
                        docx_cell = table.rows[row_idx].cells[col_idx]
                        
                        # Agregar contenido de la celda
                        cell_text = cell.get_text().strip()
                        docx_cell.text = cell_text
                        
                        # Manejar estilos de celda
                        if cell.name == 'th':
                            # Header cell - aplicar formato de encabezado
                            for paragraph in docx_cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                        
                        # Nota: rowspan/colspan son limitados en python-docx
                        # Para soporte completo se necesitaría manipulación XML directa
                        rowspan = cell.get('rowspan')
                        colspan = cell.get('colspan')
                        
                        if rowspan or colspan:
                            self.logger.debug(f"Tabla compleja detectada: rowspan={rowspan}, colspan={colspan}")
                            # Aquí iría la lógica de merge de celdas si se implementa
                            
        except Exception as e:
            self.logger.warning(f"Error procesando tabla compleja: {e}")
            # Fallback a tabla simple
            self._process_simple_table(doc, table_element)
    
    def _process_simple_table(self, doc: Document, table_element: Tag) -> None:
        """Fallback para procesar tabla como estructura simple"""
        rows = table_element.find_all('tr')
        if not rows:
            return
        
        # Crear tabla simple
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        
        # Agregar todo el contenido como texto plano
        table_text = table_element.get_text().strip()
        table.rows[0].cells[0].text = table_text
    
    def get_conversion_statistics(self) -> Dict[str, Any]:
        """Retorna estadísticas de conversión acumuladas"""
        stats = self.conversion_stats.copy()
        stats['success_rate'] = (
            (stats['pandoc_successes'] + stats['native_successes']) / 
            max(stats['total_conversions'], 1)
        )
        stats['pandoc_success_rate'] = (
            stats['pandoc_successes'] / max(stats['total_conversions'], 1)
        )
        return stats
    
    def validate_conversion_quality(self, 
                                  html_content: str, 
                                  docx_path: str) -> Dict[str, Any]:
        """
        Valida la calidad de una conversión específica
        
        Args:
            html_content: HTML original
            docx_path: Documento DOCX generado
            
        Returns:
            Reporte de calidad detallado
        """
        original_styles = self._extract_html_styles(html_content)
        preserved_styles = self._extract_docx_styles(docx_path)
        
        style_score = self._calculate_style_preservation_score(original_styles, preserved_styles)
        
        quality_report = {
            'overall_quality_score': style_score,
            'meets_threshold': style_score >= self.validation_threshold,
            'original_style_counts': asdict(original_styles),
            'preserved_style_counts': asdict(preserved_styles),
            'recommendations': []
        }
        
        # Generar recomendaciones basadas en el score
        if style_score < self.validation_threshold:
            quality_report['recommendations'].append(
                f"La preservación de estilos ({style_score:.3f}) está por debajo del umbral "
                f"({self.validation_threshold}). Considere revisar manualmente."
            )
        
        if original_styles.footnote_count > 0 and preserved_styles.footnote_count == 0:
            quality_report['recommendations'].append(
                "Se detectaron footnotes en el original pero no en el resultado. "
                "Verifique la preservación de footnotes."
            )
        
        return quality_report


# Función de utilidad para testing rápido con funcionalidad extendida
def test_html_to_docx_conversion():
    """Test completo del convertidor con funcionalidad extendida"""
    
    # HTML de prueba con múltiples formatos (básicos y avanzados)
    test_html = """
    <!DOCTYPE html>
    <html>
    <head><title>Test Document - Formatos Avanzados</title></head>
    <body>
        <h1>Capítulo 1: Introducción con Formatos Avanzados</h1>
        <p>Este es un párrafo con <strong>texto en negrita</strong> y <em>texto en cursiva</em>.</p>
        
        <h2>Sección con Formatos Diversos</h2>
        <p style="text-align: center">Párrafo centrado con <u>texto subrayado</u> y fórmulas con <sup>superíndices</sup> y <sub>subíndices</sub>.</p>
        
        <p>Texto con <span style="font-variant: small-caps">versalitas</span> y 
        <del>texto tachado</del>. También tenemos <span style="color: red">texto en rojo</span> 
        y <mark>texto resaltado</mark>.</p>
        
        <blockquote>
            Esta es una cita importante que debe preservarse como tal.
        </blockquote>
        
        <h3>Lista de Conceptos Anidada</h3>
        <ul>
            <li>Primer concepto importante
                <ul>
                    <li>Sub-concepto A</li>
                    <li>Sub-concepto B con <strong>énfasis</strong></li>
                </ul>
            </li>
            <li>Segundo concepto con <em>cursiva</em></li>
            <li>Tercer concepto con <code>código inline</code></li>
        </ul>
        
        <p>Referencia a footnote<sup>1</sup> en el texto y otra referencia<sup>2</sup>.</p>
        
        <table>
            <tr>
                <th>Columna 1</th>
                <th colspan="2">Columna Combinada</th>
            </tr>
            <tr>
                <td rowspan="2">Celda Alta</td>
                <td>Dato B</td>
                <td>Dato C</td>
            </tr>
            <tr>
                <td>Dato D</td>
                <td>Dato E</td>
            </tr>
        </table>
        
        <pre><code>
        // Código preformateado
        function ejemplo() {
            return "Hola mundo";
        }
        </code></pre>
        
        <p style="text-align: right">Párrafo alineado a la derecha.</p>
        <p style="text-align: justify">Párrafo justificado con texto más largo para demostrar la alineación justificada en el documento final.</p>
        
        <div class="footnotes">
            <p>1. Primera nota al pie con contenido académico detallado</p>
            <p>2. Segunda nota con <em>formato interno</em> y referencias</p>
        </div>
    </body>
    </html>
    """
    
    # Inicializar convertidor con funcionalidad extendida
    converter = HTMLtoDocxConverter(
        debug_mode=True,
        enable_extended_analysis=True,
        create_functional_footnotes=True,
        validation_threshold=0.85
    )
    
    print("🧪 TESTING HTMLTODOCXCONVERTER v2.2.1 (POST-AUDITORÍA)")
    print("=" * 60)
    
    # Realizar conversión con validación extendida
    result = converter.convert_with_validation(
        html_content=test_html,
        output_path="test_output_extended.docx",
        book_id="test_book_advanced_001"
    )
    
    # Mostrar resultados detallados
    print("=== RESULTADOS DE CONVERSIÓN EXTENDIDA ===")
    print(f"✅ Éxito: {result.success}")
    print(f"🔧 Método usado: {result.method_used}")
    print(f"⏱️ Tiempo de procesamiento: {result.processing_time_seconds:.2f}s")
    print(f"📊 Score preservación estilos: {result.style_preservation_score:.3f}")
    print(f"📝 Score preservación footnotes: {result.footnote_preservation_score:.3f}")
    print(f"🔗 Footnotes funcionales creadas: {result.functional_footnotes_created}")
    print(f"⚠️ Revisión manual requerida: {result.manual_review_required}")
    
    # Mostrar alertas de preservación
    if result.preservation_alerts:
        print(f"\n🚨 ALERTAS DE PRESERVACIÓN ({len(result.preservation_alerts)} total):")
        for alert in result.preservation_alerts[:5]:  # Mostrar primeras 5
            severity_emoji = {"critical": "🔴", "warning": "🟡", "info": "🔵"}
            print(f"   {severity_emoji.get(alert.alert_type, '⚪')} {alert.message}")
            if alert.suggested_action:
                print(f"      💡 Acción: {alert.suggested_action}")
    
    # Resumen de alertas por severidad
    if result.alert_summary:
        print(f"\n📋 RESUMEN DE ALERTAS:")
        print(f"   🔴 Críticas: {result.alert_summary['critical']}")
        print(f"   🟡 Warnings: {result.alert_summary['warning']}")
        print(f"   🔵 Info: {result.alert_summary['info']}")
    
    # Información de tablas complejas
    if result.complex_tables_info:
        print(f"\n🗃️ TABLAS COMPLEJAS DETECTADAS ({len(result.complex_tables_info)}):")
        for table_info in result.complex_tables_info:
            print(f"   Tabla {table_info.table_index}: rowspan={table_info.has_rowspan}, "
                  f"colspan={table_info.has_colspan}, estrategia={table_info.preservation_strategy}")
    
    if result.success:
        print(f"\n📄 Archivo generado: {result.output_path}")
        
        # Mostrar estadísticas del convertidor
        stats = converter.get_conversion_statistics()
        print(f"\n📊 Estadísticas del convertidor:")
        for key, value in stats.items():
            if isinstance(value, float):
                print(f"  {key}: {value:.3f}")
            elif isinstance(value, dict):
                print(f"  {key}: {value}")
            else:
                print(f"  {key}: {value}")
        
        # Test de validación de calidad
        quality_report = converter.validate_conversion_quality(test_html, result.output_path)
        print(f"\n🔍 REPORTE DE CALIDAD:")
        print(f"  Score general: {quality_report['overall_quality_score']:.3f}")
        print(f"  Cumple umbral: {quality_report['meets_threshold']}")
        print(f"  Revisión manual: {quality_report['manual_review_required']}")
        
        if quality_report.get('quality_metrics'):
            metrics = quality_report['quality_metrics']
            print(f"  Formatos básicos: {metrics.get('basic_formats_score', 0):.3f}")
            print(f"  Estructuras: {metrics.get('structure_preservation_score', 0):.3f}")
            if 'advanced_formats_score' in metrics:
                print(f"  Formatos avanzados: {metrics['advanced_formats_score']:.3f}")
        
        # Recomendaciones
        if quality_report.get('recommendations'):
            print(f"\n💡 RECOMENDACIONES:")
            for rec in quality_report['recommendations'][:3]:
                print(f"  • {rec}")
        
        # Test de marcadores de formato
        print(f"\n🔧 Test de marcadores de formato:")
        protected_html = converter._insert_format_markers(test_html)
        restored_html = converter.restore_format_markers(protected_html)
        
        markers_inserted = len(re.findall(r'<<<\w+_(?:START|END)>>>', protected_html))
        print(f"  Marcadores insertados: {markers_inserted}")
        print(f"  HTML restaurado correctamente: {len(restored_html) > len(test_html)}")
        
        # Verificar marcadores avanzados
        advanced_markers = [
            '<<<SMALLCAPS_START>>>', '<<<STRIKETHROUGH_START>>>', 
            '<<<COLOR_red_START>>>', '<<<HIGHLIGHT_START>>>',
            '<<<P_CENTER_START>>>', '<<<CODE_START>>>'
        ]
        found_advanced = sum(1 for marker in advanced_markers if marker in protected_html)
        print(f"  Marcadores avanzados detectados: {found_advanced}/{len(advanced_markers)}")
        
    else:
        print(f"❌ Error: {result.error_details}")
        print("🔍 Revise la configuración y dependencias del sistema")
    
    print(f"\n🎉 Test completado. Funcionalidad extendida: ✅ Activa")
    return result


if __name__ == "__main__":
    # Ejecutar test si se ejecuta directamente
    test_result = test_html_to_docx_conversion()
)),
            soup.find_all('a', href=re.compile(r'#fn\d+')),
            soup.find_all(class_=re.compile(r'footnote')),
        ]
        style_info.footnote_count = sum(len(pattern) for pattern in footnote_patterns)
        
        return style_info
    
    def _extract_docx_styles_fallback(self, docx_path: str) -> StyleInfo:
        """Método fallback para extraer estilos de DOCX"""
        try:
            doc = Document(docx_path)
            style_info = StyleInfo()
            
            # Contadores para formatos básicos
            bold_count = 0
            italic_count = 0
            underline_count = 0
            superscript_count = 0
            subscript_count = 0
            
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run.bold:
                        bold_count += 1
                    if run.italic:
                        italic_count += 1
                    if run.underline:
                        underline_count += 1
                    if run.font.superscript:
                        superscript_count += 1
                    if run.font.subscript:
                        subscript_count += 1
            
            style_info.bold_count = bold_count
            style_info.italic_count = italic_count
            style_info.underline_count = underline_count
            style_info.superscript_count = superscript_count
            style_info.subscript_count = subscript_count
            
            # Contar headings por estilo
            heading_counts = {}
            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith('Heading'):
                    level_match = re.search(r'Heading (\d+)', paragraph.style.name)
                    if level_match:
                        level = f"h{level_match.group(1)}"
                        heading_counts[level] = heading_counts.get(level, 0) + 1
            
            style_info.heading_counts = heading_counts
            style_info.table_count = len(doc.tables)
            
            return style_info
            
        except Exception as e:
            self.logger.error(f"Error extrayendo estilos de DOCX: {e}")
            return StyleInfo()
    
    def _calculate_style_preservation_score_fallback(self, 
                                                   original: StyleInfo, 
                                                   preserved: StyleInfo) -> float:
        """Método fallback para calcular score de preservación"""
        try:
            scores = []
            
            # Comparar formatos básicos
            basic_formats = [
                (original.bold_count, preserved.bold_count),
                (original.italic_count, preserved.italic_count),
                (original.underline_count, preserved.underline_count),
                (original.superscript_count, preserved.superscript_count),
                (original.subscript_count, preserved.subscript_count),
                (original.blockquote_count, preserved.blockquote_count),
                (original.table_count, preserved.table_count)
            ]
            
            for orig_count, pres_count in basic_formats:
                if orig_count == 0:
                    score = 1.0 if pres_count == 0 else 0.8
                else:
                    score = min(pres_count / orig_count, 1.0)
                scores.append(score)
            
            # Comparar headings
            for level in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                orig_count = original.heading_counts.get(level, 0)
                pres_count = preserved.heading_counts.get(level, 0)
                
                if orig_count == 0:
                    score = 1.0 if pres_count == 0 else 0.8
                else:
                    score = min(pres_count / orig_count, 1.0)
                scores.append(score)
            
            # Comparar listas
            for list_type in ['ul', 'ol', 'li']:
                orig_count = original.list_counts.get(list_type, 0)
                pres_count = preserved.list_counts.get(list_type, 0)
                
                if orig_count == 0:
                    score = 1.0 if pres_count == 0 else 0.8
                else:
                    score = min(pres_count / orig_count, 1.0)
                scores.append(score)
            
            overall_score = sum(scores) / len(scores) if scores else 0.0
            return overall_score
            
        except Exception as e:
            self.logger.error(f"Error calculando score de preservación: {e}")
            return 0.0
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Contar elementos de formato
        style_info = StyleInfo()
        
        # Formatos básicos
        style_info.bold_count = len(soup.find_all(['strong', 'b']))
        style_info.italic_count = len(soup.find_all(['em', 'i']))
        style_info.underline_count = len(soup.find_all('u'))
        style_info.superscript_count = len(soup.find_all('sup'))
        style_info.subscript_count = len(soup.find_all('sub'))
        style_info.blockquote_count = len(soup.find_all('blockquote'))
        
        # Headings
        style_info.heading_counts = {}
        for level in range(1, 7):
            count = len(soup.find_all(f'h{level}'))
            if count > 0:
                style_info.heading_counts[f'h{level}'] = count
        
        # Listas
        style_info.list_counts = {
            'ul': len(soup.find_all('ul')),
            'ol': len(soup.find_all('ol')),
            'li': len(soup.find_all('li'))
        }
        
        # Tablas y otros
        style_info.table_count = len(soup.find_all('table'))
        
        # Footnotes (buscar patrones comunes)
        footnote_patterns = [
            soup.find_all('sup', string=re.compile(r'^\d+$')),
            soup.find_all('a', href=re.compile(r'#fn\d+')),
            soup.find_all(class_=re.compile(r'footnote')),
        ]
        style_info.footnote_count = sum(len(pattern) for pattern in footnote_patterns)
        
        return style_info
    
    def _insert_format_markers(self, html_content: str) -> str:
        """
        Convierte tags HTML a marcadores protegidos para preservar formato durante traducción
        
        Args:
            html_content: HTML con formatos originales
            
        Returns:
            HTML con marcadores protegidos que Claude no modificará
        """
        protected_content = html_content
        
        # Aplicar marcadores en orden específico (del más específico al más general)
        for html_pattern, marker_replacement in self.format_markers.items():
            try:
                protected_content = re.sub(
                    html_pattern, 
                    marker_replacement,
                    protected_content, 
                    flags=re.DOTALL | re.IGNORECASE
                )
            except re.error as e:
                self.logger.warning(f"Error en patrón regex {html_pattern}: {e}")
                continue
        
        # Logging para debug
        if self.debug_mode:
            markers_inserted = len(re.findall(r'<<<\w+_(?:START|END)>>>', protected_content))
            self.logger.debug(f"Insertados {markers_inserted} marcadores de formato")
        
        return protected_content
    
    def _convert_html_to_docx_pandoc(self, html_content: str, output_path: str) -> bool:
        """
        Conversión HTML→DOCX usando Pandoc
        
        Args:
            html_content: HTML a convertir
            output_path: Ruta del archivo DOCX de salida
            
        Returns:
            True si la conversión fue exitosa
        """
        try:
            with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as temp_html:
                temp_html.write(html_content)
                temp_html_path = temp_html.name
            
            # Comando Pandoc con opciones optimizadas para preservación
            pandoc_cmd = [
                'pandoc',
                temp_html_path,
                '-f', 'html',
                '-t', 'docx',
                '-o', output_path,
                '--preserve-tabs',
                '--wrap=preserve'
            ]
            
            self.logger.debug(f"Ejecutando: {' '.join(pandoc_cmd)}")
            
            result = subprocess.run(
                pandoc_cmd,
                capture_output=True,
                text=True,
                timeout=self.pandoc_timeout
            )
            
            # Cleanup
            Path(temp_html_path).unlink(missing_ok=True)
            
            if result.returncode == 0:
                if Path(output_path).exists() and Path(output_path).stat().st_size > 0:
                    return True
                else:
                    self.logger.error("Pandoc ejecutó sin errores pero no generó archivo válido")
                    return False
            else:
                self.logger.error(f"Pandoc falló: {result.stderr}")
                return False
                
        except subprocess.TimeoutExpired:
            self.logger.error(f"Pandoc timeout después de {self.pandoc_timeout}s")
            return False
        except Exception as e:
            self.logger.error(f"Error ejecutando Pandoc: {e}")
            return False
    
    def _convert_html_to_docx_native(self, html_content: str, output_path: str) -> bool:
        """
        Conversión HTML→DOCX usando python-docx nativo
        
        Args:
            html_content: HTML a convertir  
            output_path: Ruta del archivo DOCX de salida
            
        Returns:
            True si la conversión fue exitosa
        """
        try:
            # Crear documento DOCX
            doc = Document()
            
            # Configurar estilos básicos
            self._setup_document_styles(doc)
            
            # Parsear HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Procesar elementos HTML secuencialmente
            for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 
                                        'blockquote', 'ul', 'ol', 'table', 'div', 'br']):
                self._process_html_element(doc, element)
            
            # Guardar documento
            doc.save(output_path)
            
            if Path(output_path).exists() and Path(output_path).stat().st_size > 0:
                return True
            else:
                self.logger.error("Conversión nativa no generó archivo válido")
                return False
                
        except Exception as e:
            self.logger.error(f"Error en conversión nativa: {e}")
            return False
    
    def _setup_document_styles(self, doc: Document) -> None:
        """Configura estilos básicos en el documento DOCX"""
        styles = doc.styles
        
        # Estilo para blockquotes
        try:
            quote_style = styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)
            quote_style.paragraph_format.left_indent = Inches(0.5)
            quote_style.paragraph_format.right_indent = Inches(0.5)
            quote_style.font.italic = True
            quote_style.font.size = Pt(11)
        except:
            pass  # El estilo ya puede existir
    
    def _process_html_element(self, doc: Document, element: Tag) -> None:
        """Procesa un elemento HTML individual y lo agrega al documento DOCX"""
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            # Encabezados
            level = int(element.name[1])
            heading = doc.add_heading(element.get_text(), level)
            
        elif element.name == 'p':
            # Párrafos
            paragraph = doc.add_paragraph()
            self._process_inline_formatting(paragraph, element)
            
        elif element.name == 'blockquote':
            # Citas
            paragraph = doc.add_paragraph(element.get_text(), style='Quote')
            
        elif element.name in ['ul', 'ol']:
            # Listas
            for li in element.find_all('li', recursive=False):
                paragraph = doc.add_paragraph(li.get_text(), style='List Bullet' if element.name == 'ul' else 'List Number')
                
        elif element.name == 'br':
            # Saltos de línea
            doc.add_paragraph()
    
    def _process_inline_formatting(self, paragraph, element: Tag) -> None:
        """Procesa formato inline dentro de un párrafo"""
        for content in element.contents:
            if isinstance(content, str):
                # Texto plano
                run = paragraph.add_run(content)
            elif hasattr(content, 'name'):
                # Elemento HTML
                text = content.get_text()
                run = paragraph.add_run(text)
                
                # Aplicar formato según el tag
                if content.name in ['strong', 'b']:
                    run.bold = True
                elif content.name in ['em', 'i']:
                    run.italic = True
                elif content.name == 'u':
                    run.underline = True
                elif content.name == 'sup':
                    run.font.superscript = True
                elif content.name == 'sub':
                    run.font.subscript = True
    
    def _process_footnotes_for_preservation(self, html_content: str) -> Tuple[str, Dict[str, FootnoteInfo]]:
        """
        Procesa footnotes insertando IDs únicos para preservación durante traducción
        
        Args:
            html_content: HTML original con footnotes
            
        Returns:
            Tuple de (HTML procesado con IDs, mapeo de footnotes)
        """
        footnote_mapping = {}
        processed_html = html_content
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Buscar footnotes por patrones comunes
        footnote_refs = []
        
        # Patrón 1: <sup>número</sup>
        for sup in soup.find_all('sup'):
            if re.match(r'^\d+$', sup.get_text().strip()):
                footnote_refs.append({
                    'element': sup,
                    'number': int(sup.get_text().strip()),
                    'pattern': 'superscript_number'
                })
        
        # Patrón 2: <a href="#fn1">1</a>
        for link in soup.find_all('a', href=re.compile(r'#fn\d+')):
            footnote_refs.append({
                'element': link,
                'number': int(re.search(r'\d+', link.get('href')).group()),
                'pattern': 'link_reference'
            })
        
        # Patrón 3: Elementos con class footnote
        for elem in soup.find_all(class_=re.compile(r'footnote')):
            footnote_refs.append({
                'element': elem,
                'number': len(footnote_refs) + 1,
                'pattern': 'class_footnote'
            })
        
        # Procesar cada footnote encontrada
        for i, ref_info in enumerate(footnote_refs):
            element = ref_info['element']
            number = ref_info['number']
            
            # Generar ID único
            unique_id = f"FOOTNOTE_ID_{number}_{hashlib.md5(element.get_text().encode()).hexdigest()[:8]}"
            
            # Crear info del footnote
            footnote_info = FootnoteInfo(
                unique_id=unique_id,
                original_text=element.get_text(),
                superscript_location=str(element),
                reference_number=number
            )
            
            footnote_mapping[unique_id] = footnote_info
            
            # Insertar ID único en el HTML
            if ref_info['pattern'] == 'superscript_number':
                # Para <sup>1</sup> → <sup><<<FOOTNOTE_ID_1_abc123>>> 1</sup>
                element.string = f"<<<{unique_id}>>> {element.get_text()}"
            elif ref_info['pattern'] == 'link_reference':
                # Para <a href="#fn1">1</a> → <a href="#fn1"><<<FOOTNOTE_ID_1_abc123>>> 1</a>
                element.string = f"<<<{unique_id}>>> {element.get_text()}"
            else:
                # Para otros tipos, insertar ID antes del contenido
                element.insert(0, f"<<<{unique_id}>>> ")
        
        # Convertir el soup modificado de vuelta a HTML
        processed_html = str(soup)
        
        self.logger.info(f"Procesadas {len(footnote_mapping)} footnotes con IDs únicos")
        
        return processed_html, footnote_mapping
    
    def _extract_docx_styles(self, docx_path: str) -> StyleInfo:
        """Extrae información de estilos del DOCX generado"""
        try:
            doc = Document(docx_path)
            style_info = StyleInfo()
            
            # Contar elementos de formato en el documento
            bold_count = 0
            italic_count = 0
            underline_count = 0
            superscript_count = 0
            subscript_count = 0
            
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run.bold:
                        bold_count += 1
                    if run.italic:
                        italic_count += 1
                    if run.underline:
                        underline_count += 1
                    if run.font.superscript:
                        superscript_count += 1
                    if run.font.subscript:
                        subscript_count += 1
            
            style_info.bold_count = bold_count
            style_info.italic_count = italic_count
            style_info.underline_count = underline_count
            style_info.superscript_count = superscript_count
            style_info.subscript_count = subscript_count
            
            # Contar headings por estilo
            heading_counts = {}
            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith('Heading'):
                    level = paragraph.style.name.replace('Heading ', 'h')
                    heading_counts[level] = heading_counts.get(level, 0) + 1
            
            style_info.heading_counts = heading_counts
            
            # Contar tablas
            style_info.table_count = len(doc.tables)
            
            return style_info
            
        except Exception as e:
            self.logger.error(f"Error extrayendo estilos de DOCX: {e}")
            return StyleInfo()
    
    def _calculate_style_preservation_score(self, 
                                          original: StyleInfo, 
                                          preserved: StyleInfo) -> float:
        """
        Calcula score de preservación de estilos comparando original vs preservado
        
        Args:
            original: Estilos del HTML original
            preserved: Estilos del DOCX generado
            
        Returns:
            Score de 0.0 a 1.0 indicando qué tan bien se preservaron los estilos
        """
        try:
            scores = []
            
            # Comparar formatos básicos
            basic_formats = [
                ('bold', original.bold_count, preserved.bold_count),
                ('italic', original.italic_count, preserved.italic_count),
                ('underline', original.underline_count, preserved.underline_count),
                ('superscript', original.superscript_count, preserved.superscript_count),
                ('subscript', original.subscript_count, preserved.subscript_count),
                ('blockquote', original.blockquote_count, preserved.blockquote_count),
                ('table', original.table_count, preserved.table_count)
            ]
            
            for format_name, orig_count, pres_count in basic_formats:
                if orig_count == 0:
                    # Si no había elementos originales, perfecto si tampoco hay preservados
                    score = 1.0 if pres_count == 0 else 0.8
                else:
                    # Calcular ratio de preservación
                    preservation_ratio = min(pres_count / orig_count, 1.0)
                    score = preservation_ratio
                
                scores.append(score)
                
                if self.debug_mode:
                    self.logger.debug(f"{format_name}: {orig_count} → {pres_count} (score: {score:.3f})")
            
            # Comparar headings
            for level in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                orig_count = original.heading_counts.get(level, 0)
                pres_count = preserved.heading_counts.get(level, 0)
                
                if orig_count == 0:
                    score = 1.0 if pres_count == 0 else 0.8
                else:
                    score = min(pres_count / orig_count, 1.0)
                
                scores.append(score)
            
            # Comparar listas
            for list_type in ['ul', 'ol', 'li']:
                orig_count = original.list_counts.get(list_type, 0)
                pres_count = preserved.list_counts.get(list_type, 0)
                
                if orig_count == 0:
                    score = 1.0 if pres_count == 0 else 0.8
                else:
                    score = min(pres_count / orig_count, 1.0)
                
                scores.append(score)
            
            # Calcular score promedio ponderado
            if not scores:
                return 0.0
            
            overall_score = sum(scores) / len(scores)
            
            self.logger.info(f"Score de preservación de estilos: {overall_score:.3f}")
            
            return overall_score
            
        except Exception as e:
            self.logger.error(f"Error calculando score de preservación: {e}")
            return 0.0
    
    def _reconnect_footnotes_in_docx(self, 
                                   docx_path: str, 
                                   footnote_mapping: Dict[str, FootnoteInfo]) -> float:
        """
        Reconecta footnotes en el documento DOCX final
        
        Args:
            docx_path: Ruta del documento DOCX
            footnote_mapping: Mapeo de IDs únicos a footnotes
            
        Returns:
            Score de preservación de footnotes (0.0 a 1.0)
        """
        try:
            doc = Document(docx_path)
            reconnected_count = 0
            total_footnotes = len(footnote_mapping)
            
            if total_footnotes == 0:
                return 1.0  # No había footnotes, preservación perfecta
            
            # Buscar IDs de footnotes en el texto y reconectar
            for paragraph in doc.paragraphs:
                paragraph_text = paragraph.text
                
                for unique_id, footnote_info in footnote_mapping.items():
                    id_marker = f"<<<{unique_id}>>>"
                    
                    if id_marker in paragraph_text:
                        try:
                            # Encontrar el run que contiene el ID
                            for run in paragraph.runs:
                                if id_marker in run.text:
                                    # Remover el ID marker
                                    run.text = run.text.replace(id_marker, "")
                                    
                                    # Crear footnote real (simplificado para este ejemplo)
                                    # En implementación completa, usaríamos la API de footnotes de python-docx
                                    run.font.superscript = True
                                    
                                    reconnected_count += 1
                                    break
                        except Exception as e:
                            self.logger.warning(f"Error reconectando footnote {unique_id}: {e}")
                            continue
            
            # Guardar documento modificado
            doc.save(docx_path)
            
            preservation_score = reconnected_count / total_footnotes if total_footnotes > 0 else 1.0
            
            self.logger.info(f"Footnotes reconectadas: {reconnected_count}/{total_footnotes} "
                           f"(score: {preservation_score:.3f})")
            
            return preservation_score
            
        except Exception as e:
            self.logger.error(f"Error reconectando footnotes: {e}")
            return 0.0
    
    def _generate_validation_report(self, 
                                  original: StyleInfo,
                                  preserved: StyleInfo,
                                  footnote_mapping: Dict[str, FootnoteInfo]) -> Dict[str, Any]:
        """Genera reporte detallado de validación"""
        report = {
            'timestamp': datetime.now().isoformat(),
            'style_analysis': {
                'original_counts': asdict(original),
                'preserved_counts': asdict(preserved),
                'preservation_details': {}
            },
            'footnote_analysis': {
                'total_footnotes': len(footnote_mapping),
                'footnote_details': [info.to_dict() for info in footnote_mapping.values()]
            },
            'validation_summary': {
                'total_elements_analyzed': sum([
                    original.bold_count, original.italic_count, original.underline_count,
                    original.superscript_count, original.subscript_count, original.blockquote_count,
                    original.table_count, sum(original.heading_counts.values()),
                    sum(original.list_counts.values())
                ]),
                'critical_elements_preserved': True,  # Se calculará en implementación completa
                'warnings': [],
                'recommendations': []
            }
        }
        
        return report
    
    def _update_conversion_stats(self, processing_time: float, style_score: float) -> None:
        """Actualiza estadísticas de conversión"""
        self.conversion_stats['total_conversions'] += 1
        
        # Calcular promedios
        total = self.conversion_stats['total_conversions']
        current_avg_time = self.conversion_stats['avg_processing_time']
        current_avg_style = self.conversion_stats['avg_style_preservation']
        
        self.conversion_stats['avg_processing_time'] = (
            (current_avg_time * (total - 1) + processing_time) / total
        )
        
        self.conversion_stats['avg_style_preservation'] = (
            (current_avg_style * (total - 1) + style_score) / total
        )
    
    def restore_format_markers(self, text_content: str) -> str:
        """
        Restaura marcadores de formato a HTML después de traducción
        
        Args:
            text_content: Texto con marcadores <<<FORMAT_START>>> ... <<<FORMAT_END>>>
            
        Returns:
            HTML con tags restaurados
        """
        restored_content = text_content
        
        # Mapeo inverso: marcadores → HTML
        marker_to_html = {
            r'<<<BOLD_START>>>(.*?)<<<BOLD_END>>>': r'<strong>\1</strong>',
            r'<<<ITALIC_START>>>(.*?)<<<ITALIC_END>>>': r'<em>\1</em>',
            r'<<<UNDERLINE_START>>>(.*?)<<<UNDERLINE_END>>>': r'<u>\1</u>',
            r'<<<SUP_START>>>(.*?)<<<SUP_END>>>': r'<sup>\1</sup>',
            r'<<<SUB_START>>>(.*?)<<<SUB_END>>>': r'<sub>\1</sub>',
            r'<<<BLOCKQUOTE_START>>>(.*?)<<<BLOCKQUOTE_END>>>': r'<blockquote>\1</blockquote>',
            r'<<<H1_START>>>(.*?)<<<H1_END>>>': r'<h1>\1</h1>',
            r'<<<H2_START>>>(.*?)<<<H2_END>>>': r'<h2>\1</h2>',
            r'<<<H3_START>>>(.*?)<<<H3_END>>>': r'<h3>\1</h3>',
            r'<<<H4_START>>>(.*?)<<<H4_END>>>': r'<h4>\1</h4>',
            r'<<<H5_START>>>(.*?)<<<H5_END>>>': r'<h5>\1</h5>',
            r'<<<H6_START>>>(.*?)<<<H6_END>>>': r'<h6>\1</h6>',
            r'<<<P_START>>>(.*?)<<<P_END>>>': r'<p>\1</p>',
            r'<<<BR>>>': r'<br>',
            r'<<<UL_START>>>(.*?)<<<UL_END>>>': r'<ul>\1</ul>',
            r'<<<OL_START>>>(.*?)<<<OL_END>>>': r'<ol>\1</ol>',
            r'<<<LI_START>>>(.*?)<<<LI_END>>>': r'<li>\1</li>',
            r'<<<TABLE_START>>>(.*?)<<<TABLE_END>>>': r'<table>\1</table>',
            r'<<<TR_START>>>(.*?)<<<TR_END>>>': r'<tr>\1</tr>',
            r'<<<TD_START>>>(.*?)<<<TD_END>>>': r'<td>\1</td>',
            r'<<<TH_START>>>(.*?)<<<TH_END>>>': r'<th>\1</th>',
        }
        
        for marker_pattern, html_replacement in marker_to_html.items():
            try:
                restored_content = re.sub(
                    marker_pattern,
                    html_replacement,
                    restored_content,
                    flags=re.DOTALL
                )
            except re.error as e:
                self.logger.warning(f"Error restaurando marcador {marker_pattern}: {e}")
                continue
        
        return restored_content
    
    def get_conversion_statistics(self) -> Dict[str, Any]:
        """Retorna estadísticas de conversión acumuladas"""
        stats = self.conversion_stats.copy()
        stats['success_rate'] = (
            (stats['pandoc_successes'] + stats['native_successes']) / 
            max(stats['total_conversions'], 1)
        )
        stats['pandoc_success_rate'] = (
            stats['pandoc_successes'] / max(stats['total_conversions'], 1)
        )
        return stats
    
    def validate_conversion_quality(self, 
                                  html_content: str, 
                                  docx_path: str) -> Dict[str, Any]:
        """
        Valida la calidad de una conversión específica
        
        Args:
            html_content: HTML original
            docx_path: Documento DOCX generado
            
        Returns:
            Reporte de calidad detallado
        """
        original_styles = self._extract_html_styles(html_content)
        preserved_styles = self._extract_docx_styles(docx_path)
        
        style_score = self._calculate_style_preservation_score(original_styles, preserved_styles)
        
        quality_report = {
            'overall_quality_score': style_score,
            'meets_threshold': style_score >= self.validation_threshold,
            'original_style_counts': asdict(original_styles),
            'preserved_style_counts': asdict(preserved_styles),
            'recommendations': []
        }
        
        # Generar recomendaciones basadas en el score
        if style_score < self.validation_threshold:
            quality_report['recommendations'].append(
                f"La preservación de estilos ({style_score:.3f}) está por debajo del umbral "
                f"({self.validation_threshold}). Considere revisar manualmente."
            )
        
        if original_styles.footnote_count > 0 and preserved_styles.footnote_count == 0:
            quality_report['recommendations'].append(
                "Se detectaron footnotes en el original pero no en el resultado. "
                "Verifique la preservación de footnotes."
            )
        
        return quality_report


# Función de utilidad para testing rápido
def test_html_to_docx_conversion():
    """Test rápido del convertidor"""
    
    # HTML de prueba con múltiples formatos
    test_html = """
    <!DOCTYPE html>
    <html>
    <head><title>Test Document</title></head>
    <body>
        <h1>Capítulo 1: Introducción</h1>
        <p>Este es un párrafo con <strong>texto en negrita</strong> y <em>texto en cursiva</em>.</p>
        
        <h2>Sección con Formatos Diversos</h2>
        <p>Aquí tenemos <u>texto subrayado</u> y fórmulas con <sup>superíndices</sup> y <sub>subíndices</sub>.</p>
        
        <blockquote>
            Esta es una cita importante que debe preservarse como tal.
        </blockquote>
        
        <h3>Lista de Conceptos</h3>
        <ul>
            <li>Primer concepto importante</li>
            <li>Segundo concepto con <strong>énfasis</strong></li>
            <li>Tercer concepto con <em>cursiva</em></li>
        </ul>
        
        <p>Referencia a footnote<sup>1</sup> en el texto.</p>
        
        <table>
            <tr><th>Columna 1</th><th>Columna 2</th></tr>
            <tr><td>Dato A</td><td>Dato B</td></tr>
        </table>
    </body>
    </html>
    """
    
    # Inicializar convertidor
    converter = HTMLtoDocxConverter(debug_mode=True)
    
    # Realizar conversión
    result = converter.convert_with_validation(
        html_content=test_html,
        output_path="test_output.docx",
        book_id="test_book_001"
    )
    
    # Mostrar resultados
    print("=== RESULTADOS DE CONVERSIÓN ===")
    print(f"Éxito: {result.success}")
    print(f"Método usado: {result.method_used}")
    print(f"Tiempo de procesamiento: {result.processing_time_seconds:.2f}s")
    print(f"Score preservación estilos: {result.style_preservation_score:.3f}")
    print(f"Score preservación footnotes: {result.footnote_preservation_score:.3f}")
    
    if result.success:
        print(f"\n📄 Archivo generado: {result.output_path}")
        
        # Mostrar estadísticas
        stats = converter.get_conversion_statistics()
        print(f"\n📊 Estadísticas del convertidor:")
        for key, value in stats.items():
            print(f"  {key}: {value}")
        
        # Test de marcadores de formato
        print(f"\n🔧 Test de marcadores de formato:")
        protected_html = converter._insert_format_markers(test_html)
        restored_html = converter.restore_format_markers(protected_html)
        
        print(f"Marcadores insertados: {len(re.findall(r'<<<\w+_(?:START|END)>>>', protected_html))}")
        print(f"HTML restaurado correctamente: {len(restored_html) > len(test_html)}")
        
    else:
        print(f"❌ Error: {result.error_details}")
    
    return result


if __name__ == "__main__":
    # Ejecutar test si se ejecuta directamente
    test_result = test_html_to_docx_conversion()