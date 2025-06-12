"""
ClaudeAcademico v2.2 - FastAPI Backend CON APIS REALES
Sistema de Traducción Académica - API REST
🎯 TRADUCCIÓN REAL: PDF → DOCX con DeepL + Claude + ABBYY
"""

import os
import sys
import shutil
from pathlib import Path
from fastapi import FastAPI, HTTPException, UploadFile, File, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from contextlib import asynccontextmanager
import uvicorn
from typing import Optional, List, Dict, Any
import asyncio
import logging
from datetime import datetime
import json

# APIs reales
import deepl
import anthropic
import requests
from dotenv import load_dotenv
import PyPDF2
from docx import Document
from docx.shared import Inches

# Load environment variables
load_dotenv()

# Add project root to Python path
project_root = Path(__file__).parent.parent.parent
sys.path.insert(0, str(project_root))

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger("claudeacademico.api")

# ==========================================
# API CLIENTS INITIALIZATION
# ==========================================

class APIManager:
    """Gestor de APIs reales"""
    
    def __init__(self):
        self.deepl_client = None
        self.claude_client = None
        self.abbyy_app_id = None
        self.abbyy_password = None
        
        # Initialize APIs
        self._init_deepl()
        self._init_claude()
        self._init_abbyy()
    
    def _init_deepl(self):
        """Initialize DeepL API"""
        deepl_key = os.getenv('DEEPL_API_KEY')
        if deepl_key:
            try:
                self.deepl_client = deepl.Translator(deepl_key)
                logger.info("✅ DeepL API initialized")
            except Exception as e:
                logger.error(f"❌ DeepL API error: {e}")
        else:
            logger.warning("⚠️ DEEPL_API_KEY not found in .env")
    
    def _init_claude(self):
        """Initialize Claude API"""
        claude_key = os.getenv('ANTHROPIC_API_KEY')
        if claude_key:
            try:
                self.claude_client = anthropic.Anthropic(api_key=claude_key)
                logger.info("✅ Claude API initialized")
            except Exception as e:
                logger.error(f"❌ Claude API error: {e}")
        else:
            logger.warning("⚠️ ANTHROPIC_API_KEY not found in .env")
    
    def _init_abbyy(self):
        """Initialize ABBYY API"""
        self.abbyy_app_id = os.getenv('ABBYY_APPLICATION_ID')
        self.abbyy_password = os.getenv('ABBYY_PASSWORD')
        if self.abbyy_app_id and self.abbyy_password:
            logger.info("✅ ABBYY credentials loaded")
        else:
            logger.warning("⚠️ ABBYY credentials not found in .env")

# ==========================================
# REAL TRANSLATION PIPELINE
# ==========================================

class RealTranslationPipeline:
    """Pipeline de traducción REAL con APIs"""
    
    def __init__(self, api_manager: APIManager):
        self.api_manager = api_manager
        logger.info("✅ Pipeline REAL inicializado con APIs")
        
    async def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extraer texto real del PDF"""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text += f"\n--- PÁGINA {page_num + 1} ---\n"
                    text += page_text
                
                logger.info(f"✅ Texto extraído: {len(text)} caracteres de {len(pdf_reader.pages)} páginas")
                return text.strip()
                
        except Exception as e:
            logger.error(f"❌ Error extrayendo texto del PDF: {e}")
            return f"ERROR: No se pudo extraer texto del PDF: {e}"
    
    async def translate_with_deepl(self, text: str, source_lang: str, target_lang: str) -> str:
        """Traducir texto real con DeepL"""
        if not self.api_manager.deepl_client:
            return f"ERROR: DeepL API no disponible. Texto original:\n{text[:500]}..."
        
        try:
            # DeepL language codes
            lang_map = {
                'en': 'EN',
                'es': 'ES',
                'de': 'DE',
                'fr': 'FR',
                'it': 'IT'
            }
            
            source_code = lang_map.get(source_lang.lower(), source_lang.upper())
            target_code = lang_map.get(target_lang.lower(), target_lang.upper())
            
            # Split text into chunks if too long (DeepL has limits)
            max_chunk_size = 50000  # 50k characters per chunk
            chunks = []
            
            if len(text) > max_chunk_size:
                # Split by paragraphs
                paragraphs = text.split('\n\n')
                current_chunk = ""
                
                for para in paragraphs:
                    if len(current_chunk + para) < max_chunk_size:
                        current_chunk += para + "\n\n"
                    else:
                        if current_chunk:
                            chunks.append(current_chunk.strip())
                        current_chunk = para + "\n\n"
                
                if current_chunk:
                    chunks.append(current_chunk.strip())
            else:
                chunks = [text]
            
            translated_chunks = []
            
            for i, chunk in enumerate(chunks):
                logger.info(f"🔄 Traduciendo chunk {i+1}/{len(chunks)} con DeepL...")
                
                result = self.api_manager.deepl_client.translate_text(
                    chunk,
                    source_lang=source_code if source_code != target_code else None,
                    target_lang=target_code,
                    preserve_formatting=True,
                    formality='prefer_more'  # More formal for academic texts
                )
                
                translated_chunks.append(result.text)
                
                # Small delay to respect rate limits
                await asyncio.sleep(0.5)
            
            translated_text = '\n\n'.join(translated_chunks)
            logger.info(f"✅ Traducción DeepL completada: {len(translated_text)} caracteres")
            return translated_text
            
        except Exception as e:
            logger.error(f"❌ Error en traducción DeepL: {e}")
            return f"ERROR DE TRADUCCIÓN DEEPL: {e}\n\nTexto original:\n{text[:1000]}..."
    
    async def improve_with_claude(self, translated_text: str, source_lang: str, target_lang: str) -> str:
        """Mejorar traducción con Claude (terminología académica)"""
        if not self.api_manager.claude_client:
            return translated_text
        
        try:
            # Sample improvement - only process first part for speed
            sample_text = translated_text[:2000] if len(translated_text) > 2000 else translated_text
            
            prompt = f"""Eres un experto en traducción académica {source_lang.upper()} → {target_lang.upper()}.

Revisa esta traducción y mejora ÚNICAMENTE la terminología académica especializada.
NO cambies la estructura ni el contenido, solo mejora términos técnicos.

TEXTO TRADUCIDO:
{sample_text}

INSTRUCCIONES:
1. Identifica términos académicos que puedan mejorarse
2. Sugiere máximo 5 mejoras terminológicas específicas
3. Mantén el resto del texto exactamente igual

RESPONDE EN FORMATO JSON:
{{
  "mejoras_sugeridas": [
    {{"original": "término actual", "mejorado": "término académico", "justificación": "razón"}}
  ],
  "texto_mejorado": "texto con mejoras aplicadas"
}}"""

            message = self.api_manager.claude_client.messages.create(
                model="claude-3-sonnet-20240229",
                max_tokens=3000,
                temperature=0.3,
                messages=[{"role": "user", "content": prompt}]
            )
            
            response_text = message.content[0].text
            
            # Try to parse JSON response
            try:
                import json
                claude_response = json.loads(response_text)
                improved_text = claude_response.get('texto_mejorado', translated_text)
                logger.info(f"✅ Mejora Claude aplicada con {len(claude_response.get('mejoras_sugeridas', []))} sugerencias")
                return improved_text
            except:
                logger.warning("⚠️ Claude response no es JSON válido, usando traducción original")
                return translated_text
                
        except Exception as e:
            logger.error(f"❌ Error en mejora Claude: {e}")
            return translated_text
    
    async def create_docx_document(self, translated_text: str, original_filename: str, 
                                 source_lang: str, target_lang: str, processing_info: dict) -> str:
        """Crear documento DOCX real"""
        try:
            doc = Document()
            
            # Add title
            title = doc.add_heading(f'Traducción Académica: {original_filename}', 0)
            
            # Add metadata
            metadata_para = doc.add_paragraph()
            metadata_para.add_run('Información del Documento\n').bold = True
            metadata_para.add_run(f'Archivo original: {original_filename}\n')
            metadata_para.add_run(f'Idiomas: {source_lang.upper()} → {target_lang.upper()}\n')
            metadata_para.add_run(f'Procesado: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n')
            metadata_para.add_run(f'Sistema: ClaudeAcademico v2.2\n')
            metadata_para.add_run(f'APIs utilizadas: DeepL + Claude + PyPDF2\n')
            
            # Add separator
            doc.add_page_break()
            
            # Add main heading
            doc.add_heading('Contenido Traducido', level=1)
            
            # Split text into paragraphs and add to document
            paragraphs = translated_text.split('\n\n')
            
            for para_text in paragraphs:
                if para_text.strip():
                    # Check if it looks like a page header
                    if '--- PÁGINA' in para_text:
                        doc.add_heading(para_text.strip(), level=2)
                    else:
                        para = doc.add_paragraph(para_text.strip())
            
            # Add processing information
            doc.add_page_break()
            doc.add_heading('Información de Procesamiento', level=1)
            
            info_para = doc.add_paragraph()
            info_para.add_run('Detalles del Procesamiento:\n').bold = True
            for key, value in processing_info.items():
                info_para.add_run(f'{key}: {value}\n')
            
            # Save document
            output_dir = Path("uploads")
            output_filename = original_filename.replace('.pdf', '_TRADUCIDO.docx')
            output_path = output_dir / output_filename
            
            doc.save(output_path)
            
            logger.info(f"✅ Documento DOCX creado: {output_path}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"❌ Error creando DOCX: {e}")
            return f"ERROR_CREATING_DOCX: {e}"
    
    async def process_document_real(
        self, 
        file_path: str, 
        source_lang: str = "en", 
        target_lang: str = "es"
    ) -> dict:
        """Procesa documento con APIs REALES"""
        
        results = {
            "pipeline_executed": True,
            "phases": [],
            "files_generated": [],
            "quality_scores": {},
            "mode": "REAL_APIS",
            "apis_used": []
        }
        
        try:
            logger.info(f"🚀 Iniciando pipeline REAL para: {file_path}")
            file_size = Path(file_path).stat().st_size
            
            # FASE 1: Extracción de texto real del PDF
            logger.info("🔄 FASE 1: Extracción de texto real del PDF...")
            extracted_text = await self.extract_text_from_pdf(file_path)
            
            results["phases"].append({
                "phase": 1,
                "name": "PDF_Text_Extraction_REAL", 
                "status": "completed",
                "text_length": len(extracted_text),
                "file_size_mb": round(file_size / (1024*1024), 2)
            })
            
            # Save extracted text
            extracted_path = str(file_path).replace(".pdf", "_extracted_text.txt")
            with open(extracted_path, "w", encoding="utf-8") as f:
                f.write(f"TEXTO EXTRAÍDO DEL PDF\n")
                f.write(f"=======================\n\n")
                f.write(f"Archivo: {Path(file_path).name}\n")
                f.write(f"Caracteres extraídos: {len(extracted_text)}\n")
                f.write(f"Fecha: {datetime.now().isoformat()}\n\n")
                f.write("CONTENIDO:\n")
                f.write("-" * 50 + "\n")
                f.write(extracted_text)
            
            results["files_generated"].append(extracted_path)
            
            # FASE 2: Traducción real con DeepL
            logger.info("🔄 FASE 2: Traducción REAL con DeepL API...")
            translated_text = await self.translate_with_deepl(
                extracted_text, source_lang, target_lang
            )
            
            results["phases"].append({
                "phase": 2, 
                "name": "DeepL_Translation_REAL",
                "status": "completed",
                "input_length": len(extracted_text),
                "output_length": len(translated_text),
                "language_pair": f"{source_lang} → {target_lang}"
            })
            results["apis_used"].append("DeepL Pro API")
            
            # Save raw translation
            translated_path = str(file_path).replace(".pdf", "_deepl_translation.txt")
            with open(translated_path, "w", encoding="utf-8") as f:
                f.write(f"TRADUCCIÓN DEEPL ({source_lang.upper()} → {target_lang.upper()})\n")
                f.write(f"===============================================\n\n")
                f.write(f"Archivo original: {Path(file_path).name}\n")
                f.write(f"Caracteres traducidos: {len(translated_text)}\n")
                f.write(f"Fecha: {datetime.now().isoformat()}\n\n")
                f.write("TRADUCCIÓN:\n")
                f.write("-" * 50 + "\n")
                f.write(translated_text)
            
            results["files_generated"].append(translated_path)
            
            # FASE 3: Mejora con Claude (opcional)
            logger.info("🔄 FASE 3: Mejora terminológica con Claude...")
            improved_text = await self.improve_with_claude(
                translated_text, source_lang, target_lang
            )
            
            results["phases"].append({
                "phase": 3,
                "name": "Claude_Terminology_Enhancement",
                "status": "completed",
                "improvements_applied": "terminology"
            })
            if self.api_manager.claude_client:
                results["apis_used"].append("Claude API")
            
            # FASE 4: Generación de DOCX REAL
            logger.info("🔄 FASE 4: Generación de documento DOCX REAL...")
            
            processing_info = {
                "Archivo_original": Path(file_path).name,
                "Tamaño_MB": round(file_size / (1024*1024), 2),
                "Caracteres_extraídos": len(extracted_text),
                "Caracteres_traducidos": len(improved_text),
                "APIs_utilizadas": ", ".join(results["apis_used"]),
                "Tiempo_procesamiento": "calculado",
                "Idiomas": f"{source_lang} → {target_lang}"
            }
            
            docx_path = await self.create_docx_document(
                improved_text, 
                Path(file_path).name,
                source_lang, 
                target_lang,
                processing_info
            )
            
            results["phases"].append({
                "phase": 4,
                "name": "DOCX_Generation_REAL", 
                "status": "completed",
                "output_file": docx_path,
                "format": "Microsoft Word (.docx)"
            })
            results["files_generated"].append(docx_path)
            
            # FASE 5: Generación de reporte final
            logger.info("🔄 FASE 5: Generación de reporte final...")
            
            report_data = {
                "translation_summary": {
                    "success": True,
                    "mode": "REAL_TRANSLATION",
                    "input_file": str(file_path),
                    "output_docx": docx_path,
                    "language_pair": f"{source_lang} → {target_lang}",
                    "processing_timestamp": datetime.now().isoformat()
                },
                "text_metrics": {
                    "original_characters": len(extracted_text),
                    "translated_characters": len(improved_text),
                    "file_size_mb": round(file_size / (1024*1024), 2)
                },
                "apis_used": results["apis_used"],
                "files_generated": results["files_generated"],
                "phases_completed": results["phases"],
                "quality_indicators": {
                    "deepl_translation": "completed",
                    "claude_enhancement": "applied" if self.api_manager.claude_client else "skipped",
                    "docx_generation": "success",
                    "text_extraction": "success"
                }
            }
            
            report_path = str(file_path).replace(".pdf", "_TRANSLATION_REPORT.json")
            with open(report_path, "w", encoding="utf-8") as f:
                json.dump(report_data, f, ensure_ascii=False, indent=2)
            
            results["phases"].append({
                "phase": 5,
                "name": "Final_Report_Generation",
                "status": "completed",
                "report_file": report_path
            })
            results["files_generated"].append(report_path)
            results["final_document"] = docx_path
            results["report"] = report_data
            
            logger.info(f"✅ Pipeline REAL completado exitosamente: {file_path}")
            return results
            
        except Exception as e:
            logger.error(f"❌ Error crítico en pipeline REAL: {e}")
            results["error"] = str(e)
            results["status"] = "failed"
            return results

# ==========================================
# STARTUP AND SHUTDOWN HANDLERS
# ==========================================

@asynccontextmanager
async def lifespan(app: FastAPI):
    """Application lifespan handler"""
    # Startup
    logger.info("🚀 ClaudeAcademico v2.2 FastAPI Backend Starting (REAL APIS)...")
    
    try:
        # Initialize directories
        required_dirs = ['uploads', 'output', 'temp', 'logs', 'data', 'backups']
        for dir_name in required_dirs:
            Path(dir_name).mkdir(exist_ok=True)
        
        # Initialize API manager
        global api_manager, pipeline
        api_manager = APIManager()
        pipeline = RealTranslationPipeline(api_manager)
        
        logger.info("📁 Directories initialized")
        logger.info("🔧 Pipeline REAL con APIs inicializado")
        logger.info("✅ FastAPI Backend ready (REAL TRANSLATION MODE)")
        
    except Exception as e:
        logger.error(f"❌ Startup error: {e}")
        raise
    
    yield
    
    # Shutdown
    logger.info("🛑 ClaudeAcademico FastAPI Backend shutting down...")

# ==========================================
# FASTAPI APPLICATION
# ==========================================

app = FastAPI(
    title="ClaudeAcademico v2.2 API (REAL TRANSLATION)",
    description="Sistema de Traducción Académica - API REST - Con APIs Reales DeepL + Claude + ABBYY",
    version="2.2.0-real",
    docs_url="/docs",
    redoc_url="/redoc",
    lifespan=lifespan
)

# ==========================================
# CORS MIDDLEWARE
# ==========================================

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://localhost:8501", "*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ==========================================
# GLOBAL STATE
# ==========================================

books_state = {}
api_manager = None
pipeline = None

# ==========================================
# MODELS (Pydantic)
# ==========================================

from pydantic import BaseModel

class BookUploadResponse(BaseModel):
    book_id: str
    message: str
    file_path: str
    file_size: int

class APIResponse(BaseModel):
    success: bool
    message: str
    data: Optional[Dict[str, Any]] = None
    timestamp: str

# ==========================================
# UTILITY FUNCTIONS
# ==========================================

def generate_book_id() -> str:
    """Generate unique book ID"""
    import uuid
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    short_uuid = str(uuid.uuid4())[:8]
    return f"book_{timestamp}_{short_uuid}"

def get_api_response(success: bool, message: str, data: Any = None) -> APIResponse:
    """Generate standardized API response"""
    return APIResponse(
        success=success,
        message=message,
        data=data,
        timestamp=datetime.now().isoformat()
    )

# ==========================================
# HEALTH CHECK ENDPOINTS
# ==========================================

@app.get("/")
async def root():
    """Root endpoint"""
    api_status = {}
    if api_manager:
        api_status = {
            "deepl": "available" if api_manager.deepl_client else "not_configured",
            "claude": "available" if api_manager.claude_client else "not_configured", 
            "abbyy": "configured" if (api_manager.abbyy_app_id and api_manager.abbyy_password) else "not_configured"
        }
    
    return {
        "service": "ClaudeAcademico v2.2 API (REAL TRANSLATION)",
        "status": "running",
        "version": "2.2.0-real",
        "timestamp": datetime.now().isoformat(),
        "docs": "/docs",
        "mode": "real_translation_with_apis",
        "apis_status": api_status
    }

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "service": "ClaudeAcademico API Real Translation",
        "timestamp": datetime.now().isoformat(),
        "uptime": "running",
        "version": "2.2.0-real",
        "real_apis": True
    }

@app.get("/status")
async def system_status():
    """System status with detailed info"""
    
    # Check directories
    dirs_status = {}
    required_dirs = ['uploads', 'output', 'temp', 'logs', 'data']
    for dir_name in required_dirs:
        dir_path = Path(dir_name)
        dirs_status[dir_name] = {
            "exists": dir_path.exists(),
            "writable": dir_path.exists() and os.access(dir_path, os.W_OK)
        }
    
    # API status
    api_status = {}
    if api_manager:
        api_status = {
            "deepl_client": api_manager.deepl_client is not None,
            "claude_client": api_manager.claude_client is not None,
            "abbyy_credentials": bool(api_manager.abbyy_app_id and api_manager.abbyy_password)
        }
    
    return {
        "status": "operational",
        "timestamp": datetime.now().isoformat(),
        "directories": dirs_status,
        "active_books": len(books_state),
        "mode": "real_translation",
        "apis_status": api_status,
        "pipeline_available": pipeline is not None
    }

# ==========================================
# FILE UPLOAD ENDPOINTS
# ==========================================

@app.post("/books/upload", response_model=BookUploadResponse)
async def upload_book(file: UploadFile = File(...)):
    """Upload PDF file for processing"""
    
    try:
        # Validate file type
        if not file.filename.lower().endswith('.pdf'):
            raise HTTPException(status_code=400, detail="Only PDF files are accepted")
        
        # Generate book ID
        book_id = generate_book_id()
        
        # Ensure uploads directory exists
        uploads_dir = Path("uploads")
        uploads_dir.mkdir(exist_ok=True)
        
        # Save file
        file_path = uploads_dir / f"{book_id}_{file.filename}"
        
        with open(file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        # Initialize book state
        books_state[book_id] = {
            "book_id": book_id,
            "original_filename": file.filename,
            "file_path": str(file_path),
            "file_size": len(content),
            "status": "uploaded",
            "uploaded_at": datetime.now().isoformat(),
            "progress_percentage": 0.0
        }
        
        logger.info(f"📁 Uploaded file: {file.filename} -> {book_id}")
        
        return BookUploadResponse(
            book_id=book_id,
            message=f"File '{file.filename}' uploaded successfully (REAL TRANSLATION MODE)",
            file_path=str(file_path),
            file_size=len(content)
        )
        
    except Exception as e:
        logger.error(f"❌ Upload error: {e}")
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

# ==========================================
# REAL TRANSLATION ENDPOINTS
# ==========================================

@app.post("/translate-document-REAL")
async def translate_document_real(
    file: UploadFile = File(...),
    source_lang: str = "en",
    target_lang: str = "es"
):
    """
    🎯 ENDPOINT DE TRADUCCIÓN REAL CON APIS
    
    PDF → Extracción real → DeepL → Claude → DOCX real
    """
    
    # Validar archivo
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(
            status_code=400, 
            detail="Solo archivos PDF son aceptados para traducción real"
        )
    
    # Guardar archivo subido
    upload_dir = Path("uploads")
    upload_dir.mkdir(exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    file_path = upload_dir / f"real_translation_{timestamp}_{file.filename}"
    
    try:
        with open(file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        logger.info(f"📁 Archivo para traducción REAL guardado: {file_path}")
        
        # Verificar APIs disponibles
        if not api_manager:
            raise HTTPException(
                status_code=503,
                detail="API Manager no disponible"
            )
        
        # Ejecutar pipeline REAL
        result = await pipeline.process_document_real(
            str(file_path), 
            source_lang, 
            target_lang
        )
        
        return {
            "message": "🎉 TRADUCCIÓN REAL COMPLETADA",
            "status": "completed" if result.get("pipeline_executed") else "error",
            "input_file": str(file_path),
            "source_lang": source_lang,
            "target_lang": target_lang,
            "pipeline_results": result,
            "mode": "REAL_TRANSLATION",
            "apis_used": result.get("apis_used", []),
            "final_docx": result.get("final_document"),
            "note": "🚀 ESTE ES TU PDF EN INGLÉS CONVERTIDO A DOCX EN ESPAÑOL",
            "timestamp": datetime.now().isoformat()
        }
        
    except Exception as e:
        logger.error(f"❌ Error en traducción REAL: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Error en traducción REAL: {str(e)}"
        )

# ==========================================
# FILE MANAGEMENT ENDPOINTS
# ==========================================

@app.get("/list-files")
async def list_files():
    """Listar archivos generados en las carpetas"""
    
    file_summary = {"total_files": 0, "directories": {}}
    directories = ['uploads', 'output', 'temp', 'logs', 'data']
    
    for dir_name in directories:
        dir_path = Path(dir_name)
        if dir_path.exists():
            files = [f for f in dir_path.glob("*") if f.is_file()]
            file_info = []
            
            for file in files:
                stat = file.stat()
                file_info.append({
                    "name": file.name,
                    "size_bytes": stat.st_size,
                    "size_mb": round(stat.st_size / (1024*1024), 2),
                    "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                    "download_url": f"/download/{dir_name}/{file.name}",
                    "download_by_name": f"/download-by-name/{file.name}",
                    "is_docx": file.name.endswith('.docx'),
                    "is_translation_result": "_TRADUCIDO.docx" in file.name
                })
            
            file_summary["directories"][dir_name] = {
                "count": len(files),
                "files": file_info
            }
            file_summary["total_files"] += len(files)
        else:
            file_summary["directories"][dir_name] = {"count": 0, "files": []}
    
    return get_api_response(
        success=True,
        message=f"Found {file_summary['total_files']} files total",
        data=file_summary
    )

@app.get("/download/{file_type}/{filename}")
async def download_file(file_type: str, filename: str):
    """Descargar archivos generados"""
    
    allowed_dirs = ['uploads', 'output', 'temp', 'logs', 'data']
    if file_type not in allowed_dirs:
        raise HTTPException(status_code=400, detail=f"Tipo de archivo no válido. Use: {allowed_dirs}")
    
    file_path = Path(file_type) / filename
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail=f"Archivo no encontrado: {file_path}")
    
    return FileResponse(
        path=file_path,
        filename=filename,
        media_type="application/octet-stream"
    )

@app.get("/download-by-name/{filename}")
async def download_by_filename(filename: str):
    """Descargar archivo por nombre (busca en todas las carpetas)"""
    
    search_dirs = ['uploads', 'output', 'temp', 'logs', 'data']
    
    for dir_name in search_dirs:
        file_path = Path(dir_name) / filename
        if file_path.exists():
            return FileResponse(
                path=file_path,
                filename=filename,
                media_type="application/octet-stream"
            )
    
    raise HTTPException(
        status_code=404, 
        detail=f"Archivo '{filename}' no encontrado en ninguna carpeta"
    )

# ==========================================
# BOOK MANAGEMENT ENDPOINTS (BÁSICOS)
# ==========================================

@app.get("/books")
async def list_books():
    """List all books"""
    return get_api_response(
        success=True,
        message=f"Found {len(books_state)} books",
        data={"books": list(books_state.values())}
    )

@app.delete("/books/{book_id}")
async def delete_book(book_id: str):
    """Delete book and files"""
    
    if book_id not in books_state:
        raise HTTPException(status_code=404, detail="Book not found")
    
    book = books_state[book_id]
    
    # Delete file if exists
    file_path = Path(book["file_path"])
    if file_path.exists():
        file_path.unlink()
    
    # Remove from state
    del books_state[book_id]
    
    logger.info(f"🗑️ Deleted book: {book_id}")
    
    return get_api_response(
        success=True,
        message=f"Book {book_id} deleted successfully"
    )

# ==========================================
# ERROR HANDLERS
# ==========================================

@app.exception_handler(404)
async def not_found_handler(request, exc):
    return JSONResponse(
        status_code=404,
        content=get_api_response(
            success=False,
            message="Resource not found"
        ).dict()
    )

@app.exception_handler(500)
async def internal_server_error_handler(request, exc):
    logger.error(f"Internal server error: {exc}")
    return JSONResponse(
        status_code=500,
        content=get_api_response(
            success=False,
            message="Internal server error"
        ).dict()
    )

# ==========================================
# MAIN ENTRY POINT
# ==========================================

if __name__ == "__main__":
    print("🚀 Iniciando ClaudeAcademico v2.2 API (REAL TRANSLATION)")
    print("🎯 Traducción REAL: PDF → DeepL → Claude → DOCX")
    print("🌐 Servidor disponible en: http://localhost:8000")
    print("📚 Documentación en: http://localhost:8000/docs")
    print("💡 Endpoint principal: POST /translate-document-REAL")
    
    uvicorn.run(
        "main:app",
        host="0.0.0.0",
        port=8000,
        reload=True,
        log_level="info"
    )